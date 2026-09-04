import argparse
import hashlib
import json
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

import time

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "BaoCao_KhoaHoc_GiaoTrinhDienTu_CoHocLyThuyet.docx"
FONT_NAME = "Segoe UI"

COLOR_NAVY = RGBColor(0x1F, 0x38, 0x64)
COLOR_RED = RGBColor(0xC0, 0x00, 0x00)
COLOR_GOLD = RGBColor(0xC9, 0x96, 0x3A)
COLOR_BODY = RGBColor(0x20, 0x20, 0x20)
COLOR_MUTED = RGBColor(0x59, 0x59, 0x59)


GIF_ROWS = [
    ("assets/gifs/ch1/hinh-1-06.gif", "Định lý trượt lực trên vật rắn", "images/ch1/hinh-026.png"),
    ("assets/gifs/ch1/hinh-1-09.gif", "Quy tắc hình bình hành và tổng hợp lực đồng quy", "images/ch1/hinh-033.png"),
    ("assets/gifs/ch1/hinh-1-28b.gif", "Phản lực liên kết gối tựa và ngàm phẳng", "images/ch1/hinh-118.png"),
    ("assets/gifs/ch1/hinh-1-34.gif", "Ma sát trượt trên mặt phẳng nghiêng", "images/ch1/hinh-136.png"),
    ("assets/gifs/ch1/hinh-1-35.gif", "Nón ma sát và góc ma sát tĩnh", "images/ch1/hinh-138.png"),
    ("assets/gifs/ch1/hinh-1-minh-hoa-02.gif", "Trọng tâm và tâm diện tích hình học phẳng", "images/ch1/hinh-149.png"),
    ("assets/gifs/ch2/hinh-2-07.gif", "Vận tốc và gia tốc chuyển động thẳng", "images/ch2/hinh-072.png"),
    ("assets/gifs/ch2/hinh-2-09.gif", "Gia tốc tiếp tuyến và gia tốc pháp tuyến", "images/ch2/hinh-080.png"),
    ("assets/gifs/ch2/hinh-2-15.gif", "Chuyển động quay quanh trục cố định", "images/ch2/hinh-143.png"),
    ("assets/gifs/ch2/hinh-2-16.gif", "Phân bố vận tốc trong chuyển động quay", "images/ch2/hinh-147.png"),
    ("assets/gifs/ch2/hinh-2-22.gif", "Cơ cấu truyền động bánh răng hành tinh", "images/ch2/hinh-196.png"),
    ("assets/gifs/ch2/hinh-2-26.gif", "Hợp chuyển động và gia tốc Coriolis", "images/ch2/hinh-219.png"),
    ("assets/gifs/ch2/hinh-2-34.gif", "Tâm vận tốc tức thời của vật rắn phẳng", "images/ch2/hinh-276.png"),
    ("assets/gifs/ch3/hinh-3-06.gif", "Va chạm mềm đầu đạn – toa xe cát", "images/ch3/hinh-101.png"),
    ("assets/gifs/ch3/hinh-3-10.gif", "Chuyển động trong trọng trường có cản", "images/ch3/hinh-151.png"),
    ("assets/gifs/ch3/hinh-3-11.gif", "Dao động điều hòa có cản môi trường", "images/ch3/hinh-169.png"),
    ("assets/gifs/ch3/hinh-3-17.gif", "Định lý biến thiên động lượng của chất điểm", "images/ch3/hinh-216.png"),
    ("assets/gifs/ch3/hinh-3-20.gif", "Định lý bảo toàn mô men động lượng", "images/ch3/hinh-225.png"),
    ("assets/gifs/ch3/hinh-3-21.gif", "Định lý biến thiên động năng", "images/ch3/hinh-237.png"),
    ("assets/gifs/ch3/hinh-3-22.gif", "Va chạm hai quả cầu đàn hồi", "images/ch3/hinh-244.png"),
]


SIM_INTERACTIONS = {
    "ch1-1-3": "Kéo điểm đặt và góc",
    "ch1-1-4": "Dời tâm O và tính mô men",
    "ch1-1-5": "Thu gọn hệ lực; có adapter Sim3 pilot",
    "ch1-1-6": "Khảo sát cặp lực song song",
    "ch1-2-3": "Tổng hợp hai lực đồng quy",
    "ch1-1-8": "Tách vật thể tự do",
    "ch1-3-2": "Đổi góc dây và kéo vật",
    "ch1-3-6": "Tải phân bố và tập trung",
    "ch1-5-3": "Nón ma sát; có adapter Sim3 pilot",
    "ch1-6-3": "Kéo kích thước hình ghép/khoét",
    "ch2-1-1": "Quỹ đạo, vectơ vận tốc và gia tốc",
    "ch2-1-3": "Tiếp–pháp tuyến; có adapter Sim3 pilot",
    "ch2-2-2": "Quay quanh trục; có adapter Sim3 pilot",
    "ch2-3-2": "Bánh răng–đai–puli; có adapter Sim3 pilot",
    "ch2-4-4": "Hợp chuyển động; có adapter Sim3 pilot",
    "ch2-5-2": "Giao điểm pháp tuyến vận tốc",
    "ch2-5-3": "Phân bố vận tốc; có adapter Sim3 pilot",
    "ch3-1-3": "Hệ quy chiếu; có adapter Sim3 pilot",
    "ch3-2-2": "Khảo sát lực và gia tốc",
    "ch3-2-3": "Tương tác va chạm xe",
    "ch3-3-1": "Tích phân ODE bằng RK4",
    "ch3-5-2": "Xung lực và động lượng",
    "ch3-5-3": "Mô men động lượng; có adapter Sim3 pilot",
    "ch3-5-4": "Công, động năng và thế năng",
    "ch3-6-2": "Hệ số phục hồi; có adapter Sim3 pilot",
}


def load_json(relative_path):
    path = ROOT / relative_path
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_runtime_gif_map():
    source = (ROOT / "js/gif-figures.js").read_text(encoding="utf-8")
    match = re.search(
        r"const STATIC_TO_GIF = Object\.freeze\(\{(?P<body>.*?)\}\);",
        source,
        flags=re.DOTALL,
    )
    require(match is not None, "cannot locate STATIC_TO_GIF runtime manifest")
    pairs = re.findall(
        r"'([^']+\.png)'\s*:\s*'([^']+\.gif)'",
        match.group("body"),
    )
    runtime_map = dict(pairs)
    require(len(runtime_map) == len(pairs), "duplicate static path in GIF runtime manifest")
    return runtime_map


def validate_gif_inventory(rows, runtime_map, release_gif_paths):
    report_map = {fallback: gif for gif, _, fallback in rows}
    require(len(report_map) == len(rows), "duplicate PNG fallback in report GIF inventory")
    require(report_map == runtime_map, "report GIF inventory differs from runtime manifest")
    require(
        set(report_map.values()) == set(release_gif_paths),
        "runtime GIF inventory differs from candidate release manifest",
    )
    return report_map


def require(condition, message):
    if not condition:
        raise ValueError(message)


def sha256_file(path):
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def evidence_record(registry, gate_id):
    matches = [record for record in registry["records"] if record["gateId"] == gate_id]
    require(len(matches) == 1, f"expected one evidence record for {gate_id}")
    return matches[0]


def evidence_input_hashes(record, root=ROOT):
    log_path = Path(root) / record["artifact"]
    require(log_path.is_file(), f"missing evidence log: {record['artifact']}")
    expected_artifact_hash = record.get("hash", "")
    require(
        expected_artifact_hash.startswith("sha256:")
        and len(expected_artifact_hash) == len("sha256:") + 64,
        f"invalid evidence artifact hash: {record['artifact']}",
    )
    require(
        sha256_file(log_path) == expected_artifact_hash.removeprefix("sha256:"),
        f"stale evidence log: {record['artifact']}",
    )
    hashes = {}
    pattern = re.compile(r"^sha256:([a-f0-9]{64})\s+(.+)$")
    for line in log_path.read_text(encoding="utf-8-sig").splitlines():
        match = pattern.match(line.strip())
        if match:
            hashes[match.group(2).replace("\\", "/")] = match.group(1)
    return hashes


def verify_hashed_input(relative_path, expected_hash, label):
    path = ROOT / relative_path
    require(path.is_file(), f"missing {label}: {relative_path}")
    actual_hash = sha256_file(path)
    require(actual_hash == expected_hash, f"stale {label}: {relative_path}")
    return actual_hash


def candidate_evidence_binding(acceptance, record, input_hashes, expected_hashes):
    acceptance_matches = [
        gate
        for gate in acceptance["gates"]
        if gate["gateId"] == "release-candidate-inventory"
    ]
    join_mismatches = []
    if len(acceptance_matches) != 1:
        join_mismatches.append("acceptance gate count")
    else:
        gate = acceptance_matches[0]
        for field in ("status", "artifact", "hash", "observedAt"):
            if gate.get(field) != record.get(field):
                join_mismatches.append(field)

    missing_inputs = sorted(set(expected_hashes) - set(input_hashes))
    hash_mismatches = sorted(
        path
        for path, expected_hash in expected_hashes.items()
        if path in input_hashes and input_hashes[path] != expected_hash
    )
    current = (
        record.get("status") == "pass"
        and not join_mismatches
        and not missing_inputs
        and not hash_mismatches
    )
    return {
        "current": current,
        "missingInputs": missing_inputs,
        "hashMismatches": hash_mismatches,
        "joinMismatches": join_mismatches,
    }


def decision_profile(acceptance, candidate_evidence_current):
    decision = acceptance["releaseDecision"]["decision"]
    require(decision in {"approved", "rejected", "blocked"}, "unknown release decision")
    if not candidate_evidence_current:
        return {
            "key": "evidence-mismatch",
            "accepted": False,
            "cover": "DỰ THẢO — BẰNG CHỨNG KHÔNG CÙNG SNAPSHOT",
            "header": "DỰ THẢO — CẦN LÀM MỚI BẰNG CHỨNG CANDIDATE",
            "summary": (
                "Snapshot nghiệm thu và candidate hiện hành chưa được ràng buộc bởi cùng bộ "
                "hash đầu vào; không được quy kết trạng thái gate cũ cho candidate mới."
            ),
            "conclusion": (
                "Báo cáo chỉ được dùng để nhận diện chênh lệch bằng chứng. Cần chạy lại cổng "
                "candidate và acceptance trên cùng phiên bản trước mọi quyết định phát hành."
            ),
        }
    if decision == "approved":
        return {
            "key": "approved",
            "accepted": True,
            "cover": "TRẠNG THÁI: ĐÃ ĐỦ CỔNG BẰNG CHỨNG",
            "header": "BÁO CÁO KỸ THUẬT — ĐỦ CỔNG BẰNG CHỨNG",
            "summary": "Mọi cổng bắt buộc trong snapshot candidate hiện hành đã pass.",
            "conclusion": (
                "Ma trận bằng chứng hiện hành đã đạt toàn bộ cổng bắt buộc. Tài liệu có thể "
                "chuyển sang quyết định phát hành theo thẩm quyền của cơ sở."
            ),
        }
    if decision == "rejected":
        return {
            "key": "rejected",
            "accepted": False,
            "cover": "BỊ TỪ CHỐI — CÓ CỔNG KIỂM TRA THẤT BẠI",
            "header": "BÁO CÁO KỸ THUẬT — CANDIDATE BỊ TỪ CHỐI",
            "summary": (
                "Ít nhất một cổng bắt buộc đã fail; candidate không đủ điều kiện chuyển sang "
                "phát hành hoặc nghiệm thu."
            ),
            "conclusion": (
                "Candidate bị từ chối do có cổng bắt buộc thất bại. Phải sửa nguyên nhân, "
                "tạo bằng chứng mới và chạy lại acceptance trước khi xem xét tiếp."
            ),
        }
    return {
        "key": "blocked",
        "accepted": False,
        "cover": "DỰ THẢO — 20 CỔNG ĐẠT, CHỜ 4 ĐÁNH GIÁ ĐỘC LẬP",
        "header": "DỰ THẢO — CHỜ NGHIỆM THU ĐỘC LẬP",
        "summary": (
            "Hai mươi cổng kỹ thuật đã pass, không có cổng fail; bốn đánh giá độc lập "
            "chưa hoàn tất nên candidate chưa đủ điều kiện phát hành chính thức."
        ),
        "conclusion": (
            "Candidate đã có hiện vật kỹ thuật khóa hash và 20 cổng pass nhưng vẫn thiếu điều kiện "
            "bắt buộc độc lập. Hồ sơ phù hợp để Hội đồng xem xét theo hướng thông qua có điều kiện; "
            "chưa có cơ sở để tuyên bố nghiệm thu học thuật, tuân thủ WCAG toàn hệ thống, tương thích "
            "LMS thực tế hoặc đưa vào giảng dạy chính thức trước khi bốn đánh giá độc lập được đóng."
        ),
    }


def normalize_docx_package(path):
    path = Path(path)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source:
        entries = [(name, source.read(name)) for name in sorted(source.namelist())]
    with zipfile.ZipFile(
        temporary,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for name, payload in entries:
            info = zipfile.ZipInfo(name, (1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, payload)
    last_error = None
    for attempt in range(10):
        try:
            os.replace(temporary, path)
            return
        except OSError as error:
            last_error = error
            time.sleep(0.2 * (attempt + 1))
    if last_error:
        raise last_error


def load_evidence():
    acceptance = load_json("data/acceptance-report.json")
    release_candidate = load_json("data/release-candidate.json")
    release_summary = load_json(release_candidate["summaryPath"])
    release_root = Path(release_candidate["summaryPath"]).parent
    release_manifest_path = (
        release_root / release_summary["manifest"]["path"]
    ).as_posix()
    release_manifest = load_json(release_manifest_path)
    evidence_registry = load_json("data/evidence-registry.json")
    accessibility = load_json("data/accessibility-baseline.json")
    lms_targets = load_json("data/lms-targets.json")
    legal = load_json("data/legal-standards-register.json")
    learning_outcomes = load_json("data/learning-outcomes.json")
    content_manifest = load_json("data/content-manifest.json")
    equations = load_json("data/equation_mapping.json")
    quizzes = {
        chapter: load_json(f"data/quiz-ch{chapter}.json").get("items", [])
        for chapter in (1, 2, 3)
    }
    simulation_specs = load_json("data/simulation-specifications.json")
    sim3_reviews = load_json("data/sim3-pedagogical-reviews.json")

    gate_summary = acceptance["gateSummary"]
    gates = acceptance["gates"]
    status_total = sum(
        gate_summary[key] for key in ("pass", "fail", "blocked", "notRun")
    )
    if gate_summary["fail"]:
        expected_overall = "fail"
        expected_decision = "rejected"
    elif gate_summary["blocked"] or gate_summary["notRun"]:
        expected_overall = "blocked"
        expected_decision = "blocked"
    else:
        expected_overall = "pass"
        expected_decision = "approved"
    require(
        acceptance["overallStatus"] == expected_overall,
        "acceptance overall status contradicts gate summary",
    )
    require(
        acceptance["releaseDecision"]["decision"] == expected_decision,
        "release decision contradicts gate summary",
    )
    require(
        release_summary["releaseVersion"] == release_candidate["releaseVersion"],
        "release version mismatch",
    )
    require(
        release_summary["package"]["sha256"] == release_candidate["packageSha256"],
        "release SHA-256 mismatch",
    )
    verify_hashed_input(
        release_manifest_path,
        release_summary["manifest"]["sha256"],
        "release manifest",
    )

    release_files = {record["path"]: record for record in release_manifest["files"]}
    content_provenance = release_manifest["provenance"]["contentManifest"]
    verify_hashed_input(
        content_provenance["path"],
        content_provenance["sha256"],
        "candidate content manifest",
    )
    for chapter in (1, 2, 3):
        quiz_path = f"data/quiz-ch{chapter}.json"
        require(quiz_path in release_files, f"candidate omits {quiz_path}")
        verify_hashed_input(
            quiz_path,
            release_files[quiz_path]["sha256"],
            "candidate quiz",
        )

    candidate_record = evidence_record(evidence_registry, "release-candidate-inventory")
    candidate_hashes = evidence_input_hashes(candidate_record)
    package_path = (release_root / release_summary["package"]["path"]).as_posix()
    verify_hashed_input(
        package_path,
        release_summary["package"]["sha256"],
        "candidate package",
    )
    expected_candidate_hashes = {
        release_candidate["summaryPath"]: sha256_file(
            ROOT / release_candidate["summaryPath"]
        ),
        release_manifest_path: release_summary["manifest"]["sha256"],
        package_path: release_summary["package"]["sha256"],
    }
    candidate_binding = candidate_evidence_binding(
        acceptance,
        candidate_record,
        candidate_hashes,
        expected_candidate_hashes,
    )

    equation_record = evidence_record(evidence_registry, "equations")
    equation_hashes = evidence_input_hashes(equation_record)
    require(
        "data/equation_mapping.json" in equation_hashes,
        "equation evidence omits equation mapping",
    )
    verify_hashed_input(
        "data/equation_mapping.json",
        equation_hashes["data/equation_mapping.json"],
        "equation evidence input",
    )

    simulation_record = evidence_record(
        evidence_registry,
        "simulation-evidence-currentness",
    )
    simulation_hashes = evidence_input_hashes(simulation_record)
    simulation_inputs = {
        path.replace("\\", "/") for path in simulation_record["inputs"]
    }
    for relative_path in (
        "data/simulation-specifications.json",
        "data/sim3-pedagogical-reviews.json",
    ):
        require(relative_path in simulation_hashes, f"simulation evidence omits {relative_path}")
        verify_hashed_input(
            relative_path,
            simulation_hashes[relative_path],
            "simulation evidence input",
        )

    image_provenance = {}
    candidate_images = (
        "images/ch1/hinh-002.png",
        "images/ch1/hinh-078.png",
        "images/ch2/hinh-196.png",
    )
    for relative_path in candidate_images:
        require(relative_path in release_files, f"candidate omits image: {relative_path}")
        digest = verify_hashed_input(
            relative_path,
            release_files[relative_path]["sha256"],
            "candidate image",
        )
        image_provenance[relative_path] = {
            "sha256": digest,
            "authority": f"candidate manifest {release_summary['releaseVersion']}",
        }

    simulation_images = (
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch1-6-3-negative-area-win32.png",
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch2-4-4-coriolis-callout-win32.png",
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch2-3-2-transmission-win32.png",
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch3-6-2-collision-after-win32.png",
    )
    for relative_path in simulation_images:
        require(relative_path in simulation_inputs, f"simulation evidence omits {relative_path}")
        require(relative_path in simulation_hashes, f"simulation log omits {relative_path}")
        digest = verify_hashed_input(
            relative_path,
            simulation_hashes[relative_path],
            "simulation screenshot",
        )
        image_provenance[relative_path] = {
            "sha256": digest,
            "authority": "simulation-evidence-currentness",
        }

    presentation_images = (
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-01-trang-chu-desktop-1440x1000.png",
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-02-trang-chu-mobile-390x844.png",
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-04-mo-men-ch1-1-4-1440x1000.png",
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-06-pdf-viewer-1440x1000.png",
    )
    for relative_path in presentation_images:
        capture_path = ROOT / relative_path
        require(capture_path.is_file(), f"missing presentation capture: {relative_path}")
        image_provenance[relative_path] = {
            "sha256": sha256_file(capture_path),
            "authority": "browser capture bound to presentation evidence; not independent review",
        }

    runtime_manifest_path = "js/gif-figures.js"
    require(
        runtime_manifest_path in release_files,
        "candidate omits GIF runtime manifest",
    )
    verify_hashed_input(
        runtime_manifest_path,
        release_files[runtime_manifest_path]["sha256"],
        "candidate GIF runtime manifest",
    )
    runtime_gif_map = load_runtime_gif_map()
    release_gif_paths = {
        path
        for path in release_files
        if path.startswith("assets/gifs/") and path.endswith(".gif")
    }
    validate_gif_inventory(GIF_ROWS, runtime_gif_map, release_gif_paths)

    for gif_path, _, fallback_path in GIF_ROWS:
        for relative_path in (gif_path, fallback_path):
            require(relative_path in release_files, f"candidate omits media: {relative_path}")
            verify_hashed_input(
                relative_path,
                release_files[relative_path]["sha256"],
                "candidate media",
            )

    routes = content_manifest["routes"]
    chapter_counts = {
        chapter: sum(
            route.get("routeId", "").startswith(f"ch{chapter}")
            for route in routes
        )
        for chapter in (1, 2, 3)
    }
    other_routes = len(routes) - sum(chapter_counts.values())
    quiz_counts = {chapter: len(items) for chapter, items in quizzes.items()}
    specs = simulation_specs["specifications"]
    sim3_records = sim3_reviews["reviews"]
    require(other_routes >= 0, "chapter route counts exceed total routes")
    require(specs, "simulation specifications are empty")
    require(set(SIM_INTERACTIONS) == {spec["id"] for spec in specs}, "simulation interaction map drift")

    return {
        "acceptance": acceptance,
        "releaseCandidate": release_candidate,
        "releaseSummary": release_summary,
        "releaseManifest": release_manifest,
        "accessibility": accessibility,
        "lmsTargets": lms_targets,
        "legal": legal,
        "learningOutcomes": learning_outcomes,
        "routes": routes,
        "chapterCounts": chapter_counts,
        "otherRoutes": other_routes,
        "equationCount": len(equations),
        "equationOccurrenceCount": sum(
            len(route.get("equationRefs", [])) for route in routes
        ),
        "quizCounts": quiz_counts,
        "simulationSpecs": specs,
        "sim3Reviews": sim3_records,
        "imageProvenance": image_provenance,
        "candidateEvidence": {
            **candidate_binding,
            "observedAt": candidate_record["observedAt"],
        },
    }


def set_cell_background(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    tc_mar = tc_pr.find(qn("w:tcMar"))
    v_align = tc_pr.find(qn("w:vAlign"))
    if tc_mar is not None:
        tc_mar.addprevious(shd)
    elif v_align is not None:
        v_align.addprevious(shd)
    else:
        tc_pr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        "</w:tcMar>"
    )
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is not None:
        v_align.addprevious(tc_mar)
    else:
        tc_pr.append(tc_mar)


def set_cell_border(cell, color="D0D0D0", size=2):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        "</w:tcBorders>"
    )
    shd = tc_pr.find(qn("w:shd"))
    tc_mar = tc_pr.find(qn("w:tcMar"))
    v_align = tc_pr.find(qn("w:vAlign"))
    if shd is not None:
        shd.addprevious(borders)
    elif tc_mar is not None:
        tc_mar.addprevious(borders)
    elif v_align is not None:
        v_align.addprevious(borders)
    else:
        tc_pr.append(borders)


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    tr_pr.append(header)

def set_run_font(run, size=10, color=COLOR_BODY, bold=False, italic=False):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    specs = {
        "Normal": (10.5, COLOR_BODY, False),
        "Title": (20, COLOR_NAVY, True),
        "Heading 1": (15, COLOR_RED, True),
        "Heading 2": (12.5, COLOR_NAVY, True),
        "Heading 3": (11, COLOR_BODY, True),
        "Caption": (9, COLOR_MUTED, False),
    }
    for name, (size, color, bold) in specs.items():
        style = styles[name]
        style.font.name = FONT_NAME
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)
    styles["Heading 3"].paragraph_format.keep_with_next = True
    styles["Caption"].paragraph_format.space_before = Pt(2)
    styles["Caption"].paragraph_format.space_after = Pt(8)
    styles["Caption"].paragraph_format.keep_with_next = True


def add_field(paragraph, instruction, cached_value=""):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    value_run = paragraph.add_run(cached_value)
    set_run_font(value_run, size=9, color=COLOR_MUTED)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def enable_field_updates(doc):
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        settings.remove(existing)
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is not None:
        compat.addprevious(uf)
    else:
        settings.append(uf)

def add_caption(doc, label, sequence_name, text):
    counts = getattr(doc, "_report_sequence_counts", {})
    counts[sequence_name] = counts.get(sequence_name, 0) + 1
    doc._report_sequence_counts = counts
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run(f"{label} ")
    set_run_font(prefix, size=9, color=COLOR_MUTED, italic=True)
    add_field(
        paragraph,
        f" SEQ {sequence_name} \\* ARABIC ",
        str(counts[sequence_name]),
    )
    suffix = paragraph.add_run(f". {text}")
    set_run_font(suffix, size=9, color=COLOR_MUTED, italic=True)
    return paragraph

def set_image_alt(inline_shape, alt_text):
    properties = inline_shape._inline.docPr
    properties.set("title", alt_text[:120])
    properties.set("descr", alt_text)


def add_callout(doc, items, title, border_color="1F3864", fill_color="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.0)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        '<w:bottom w:val="none"/>'
        '<w:right w:val="none"/>'
        "</w:tcBorders>"
    )
    tc_pr.append(borders)
    set_cell_background(cell, fill_color)
    set_cell_margins(cell, top=160, bottom=160, left=240, right=200)
    heading = cell.paragraphs[0]
    heading.paragraph_format.space_after = Pt(4)
    title_run = heading.add_run(title)
    set_run_font(title_run, size=10.5, color=COLOR_NAVY, bold=True)
    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, size=9.5)


def add_block_diagram(doc, blocks, caption):
    require(blocks, "block diagram requires at least one block")
    arrow_width_cm = 0.45
    box_width_cm = (16.0 - arrow_width_cm * (len(blocks) - 1)) / len(blocks)
    columns = len(blocks) * 2 - 1
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, block in enumerate(blocks):
        cell = table.cell(0, index * 2)
        cell.width = Cm(box_width_cm)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_background(cell, "E8EEF7" if index % 2 == 0 else "F7EEDC")
        set_cell_border(cell, color="8EA9C1", size=6)
        set_cell_margins(cell, top=150, bottom=150, left=100, right=100)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(block)
        set_run_font(run, size=8.5, color=COLOR_NAVY, bold=True)
        if index < len(blocks) - 1:
            arrow_cell = table.cell(0, index * 2 + 1)
            arrow_cell.width = Cm(arrow_width_cm)
            arrow_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            arrow = arrow_cell.paragraphs[0]
            arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_run = arrow.add_run("→")
            set_run_font(arrow_run, size=14, color=COLOR_RED, bold=True)
    add_caption(doc, "Sơ đồ", "Diagram", caption)


def add_data_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_row = table.rows[0]
    mark_repeat_header(header_row)
    for index, heading in enumerate(headers):
        cell = header_row.cells[index]
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_background(cell, "1F3864")
        set_cell_border(cell, color="FFFFFF", size=2)
        set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(heading))
        set_run_font(run, size=9, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.width = widths[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_background(cell, "F4F6F9" if row_index % 2 else "FFFFFF")
            set_cell_border(cell)
            set_cell_margins(cell, top=65, bottom=65, left=80, right=80)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column_index == 0 and len(headers) > 2
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
    return table


def format_megabytes(size_bytes):
    return f"{size_bytes / 1_000_000:.1f} MB ({size_bytes:,} byte; {size_bytes / 1_048_576:.1f} MiB)".replace(",", ".")


def build_report(output_path=DEFAULT_OUTPUT, evidence=None):
    evidence = load_evidence() if evidence is None else evidence
    acceptance = evidence["acceptance"]
    gate_summary = acceptance["gateSummary"]
    release = evidence["releaseSummary"]
    release_candidate = evidence["releaseCandidate"]
    accessibility = evidence["accessibility"]
    lms_targets = evidence["lmsTargets"]
    profile = decision_profile(
        acceptance,
        evidence["candidateEvidence"]["current"],
    )
    gate_status = {
        gate["gateId"]: gate["status"] for gate in acceptance["gates"]
    }

    output_path = Path(output_path)
    doc = Document()
    configure_styles(doc)
    enable_field_updates(doc)
    snapshot_time = datetime.fromisoformat(
        acceptance["generatedAt"].replace("Z", "+00:00")
    ).replace(tzinfo=None)
    doc.core_properties.title = "Báo cáo khoa học Giáo trình điện tử Cơ học lý thuyết"
    doc.core_properties.subject = "Báo cáo kỹ thuật phục vụ thẩm định độc lập"
    doc.core_properties.author = "Dự án Giáo trình điện tử Cơ học lý thuyết"
    doc.core_properties.created = snapshot_time
    doc.core_properties.modified = snapshot_time

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.different_first_page_header_footer = True

        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""

        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_text = profile["header"]
        header_run = header_paragraph.add_run(header_text)
        set_run_font(header_run, size=8.5, color=COLOR_MUTED, italic=True)

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_paragraph.add_run("Báo cáo kỹ thuật có kiểm soát bằng chứng • Trang ")
        set_run_font(footer_run, size=8.5, color=COLOR_MUTED)
        add_field(footer_paragraph, " PAGE ", "1")
        middle_run = footer_paragraph.add_run(" / ")
        set_run_font(middle_run, size=8.5, color=COLOR_MUTED)
        add_field(footer_paragraph, " NUMPAGES ", "1")

    def add_h1(text, new_page=False):
        paragraph = doc.add_paragraph(style="Heading 1")
        paragraph.paragraph_format.page_break_before = new_page
        run = paragraph.add_run(text.upper())
        set_run_font(run, size=15, color=COLOR_RED, bold=True)
        p_pr = paragraph._p.get_or_add_pPr()
        p_pr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="12" w:space="4" w:color="C00000"/>'
                "</w:pBdr>"
            )
        )
        return paragraph

    def add_h2(text):
        paragraph = doc.add_paragraph(style="Heading 2")
        run = paragraph.add_run(text)
        set_run_font(run, size=12.5, color=COLOR_NAVY, bold=True)
        return paragraph


    def add_body(text, bold_prefix=None, italic=False):
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.75)
        if bold_prefix:
            prefix = paragraph.add_run(bold_prefix)
            set_run_font(prefix, size=10.5, color=COLOR_NAVY, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, italic=italic)
        return paragraph

    def add_bullet(text, bold_prefix=None):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        if bold_prefix:
            prefix = paragraph.add_run(bold_prefix)
            set_run_font(prefix, size=10, color=COLOR_NAVY, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, size=10)
        return paragraph

    def add_image_box(relative_path, caption, alt_text, width_cm=13.5):
        image_path = ROOT / relative_path
        provenance = evidence["imageProvenance"].get(relative_path)
        require(provenance is not None, f"unbound report image: {relative_path}")
        require(image_path.is_file(), f"missing report image: {relative_path}")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(2)
        shape = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
        set_image_alt(shape, alt_text)
        add_caption(
            doc,
            "Hình",
            "Figure",
            f"{caption} Nguồn: {relative_path}; SHA-256 "
            f"{provenance['sha256']}; thẩm quyền: {provenance['authority']}.",
        )

    def add_meta_line(label, value):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10, color=COLOR_NAVY, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10)

    # Neutral project cover; institutional mastheads require separate authority evidence.
    cover_line = doc.add_paragraph()
    cover_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_line.paragraph_format.space_before = Pt(10)
    cover_run = cover_line.add_run(
        "DỰ ÁN GIÁO TRÌNH ĐIỆN TỬ CƠ HỌC LÝ THUYẾT"
    )
    set_run_font(cover_run, size=11, color=COLOR_NAVY, bold=True)

    council_line = doc.add_paragraph()
    council_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    council_line.paragraph_format.space_after = Pt(26)
    council_run = council_line.add_run(
        "BÁO CÁO KỸ THUẬT PHỤC VỤ THẨM ĐỊNH ĐỘC LẬP"
    )
    set_run_font(council_run, size=12, color=COLOR_RED, bold=True)

    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    title_cell.width = Cm(16.0)
    set_cell_background(title_cell, "1F3864")
    set_cell_margins(title_cell, top=280, bottom=280, left=240, right=240)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(
        "BÁO CÁO KHOA HỌC\n"
        "ĐÁNH GIÁ TOÀN DIỆN QUY CÁCH TRÌNH BÀY, KẾT QUẢ, CHUẨN TIÊU CHÍ "
        "VÀ PHƯƠNG PHÁP LUẬN XÂY DỰNG MÔ PHỎNG"
    )
    set_run_font(title_run, size=15.5, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    subtitle = title_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    subtitle_run = subtitle.add_run(
        "GIÁO TRÌNH ĐIỆN TỬ CƠ HỌC LÝ THUYẾT "
        "(TĨNH HỌC – ĐỘNG HỌC – ĐỘNG LỰC HỌC)"
    )
    set_run_font(subtitle_run, size=12, color=RGBColor(0xDB, 0xB3, 0x6A), bold=True)

    status_table = doc.add_table(rows=1, cols=1)
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_cell = status_table.cell(0, 0)
    set_cell_background(status_cell, "FCE4D6" if not profile["accepted"] else "E2F0D9")
    set_cell_border(
        status_cell,
        color="C00000" if not profile["accepted"] else "548235",
        size=10,
    )
    status_paragraph = status_cell.paragraphs[0]
    status_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status_paragraph.add_run(profile["cover"])
    set_run_font(
        status_run,
        size=11,
        color=COLOR_RED if not profile["accepted"] else COLOR_NAVY,
        bold=True,
    )

    add_meta_line("Đối tượng thẩm định", "Giáo trình điện tử Cơ học lý thuyết chạy tĩnh trên trình duyệt")
    add_meta_line("Nguồn nội dung chuẩn", "CoHocLyThuyet_Full_New.docx và CoHocLyThuyet.pdf")
    add_meta_line(
        "Phiên bản candidate",
        f"{release['releaseVersion']} — {release['staging']['fileCount']} tệp — "
        f"SHA-256 {release['package']['sha256']}",
    )
    add_meta_line(
        "Snapshot nghiệm thu",
        f"{acceptance['generatedAt']} — overallStatus={acceptance['overallStatus']} — "
        f"decision={acceptance['releaseDecision']['decision']}",
    )
    binding = evidence["candidateEvidence"]
    binding_issues = []
    if binding["missingInputs"]:
        binding_issues.append("thiếu input: " + ", ".join(binding["missingInputs"]))
    if binding["hashMismatches"]:
        binding_issues.append(
            "hash lệch: " + ", ".join(binding["hashMismatches"])
        )
    if binding["joinMismatches"]:
        binding_issues.append(
            "acceptance join lệch: " + ", ".join(binding["joinMismatches"])
        )
    add_meta_line(
        "Ràng buộc snapshot–candidate",
        (
            f"current; observedAt={binding['observedAt']}"
            if binding["current"]
            else "stale; " + "; ".join(binding_issues)
        ),
    )
    add_meta_line(
        "Tổng hợp cổng trong snapshot",
        f"{gate_summary['pass']}/{gate_summary['total']} pass; "
        f"{gate_summary['fail']} fail; {gate_summary['blocked']} blocked; "
        f"{gate_summary['notRun']} not run",
    )
    add_meta_line(
        "Giới hạn tuyên bố",
        "Không phải chứng nhận học thuật, WCAG, CDIO/ABET hoặc bằng chứng nhập LMS.",
    )

    doc.add_page_break()

    # TOC
    toc_title = doc.add_paragraph(style="Title")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("MỤC LỤC")
    set_run_font(toc_run, size=20, color=COLOR_NAVY, bold=True)
    toc_paragraph = doc.add_paragraph()
    add_field(
        toc_paragraph,
        ' TOC \\o "1-3" \\h \\z \\u ',
        "Mục lục sẽ được cập nhật khi mở tài liệu trong Microsoft Word.",
    )
    doc.add_page_break()

    # Executive summary and method
    add_h1("Tóm tắt điều hành")
    add_body(profile["summary"])
    add_callout(
        doc,
        [
            f"Snapshot {acceptance['generatedAt']}: {gate_summary['pass']} pass; "
            f"{gate_summary['fail']} fail; {gate_summary['blocked']} blocked; "
            f"{gate_summary['notRun']} chưa chạy.",
            (
                "Candidate và snapshot acceptance đã cùng bộ input hash."
                if binding["current"]
                else "Candidate hiện hành chưa được cổng acceptance snapshot này xác nhận."
            ),
            "Kết quả tự động chứng minh hợp đồng kỹ thuật trong phạm vi đã khai báo; "
            "không thay thế kết luận của chuyên gia.",
            "Mọi tuyên bố về WCAG, CDIO/ABET, pháp lý và khả năng nhập LMS được giữ ở mức "
            "sơ bộ cho đến khi có bằng chứng độc lập.",
            "Sim2 là lớp canonical; Sim3 là pilot tùy chọn và phải fallback về Sim2.",
        ],
        "KẾT LUẬN ĐIỀU HÀNH DỰA TRÊN BẰNG CHỨNG",
        border_color="548235" if profile["accepted"] else "C00000",
        fill_color="F2F8EF" if profile["accepted"] else "FFF4F0",
    )

    add_h2("Mục tiêu, phạm vi và phương pháp")
    add_body(
        "Báo cáo đánh giá quy cách trình bày, kiến trúc kỹ thuật, phạm vi mô phỏng, "
        "khả năng tiếp cận, gói LMS và mức độ ràng buộc giữa candidate với snapshot bằng chứng."
    )
    add_bullet(
        "Đếm route, công thức, câu hỏi và mô phỏng trực tiếp từ manifest hoặc registry chuẩn.",
        "Định lượng: ",
    )
    add_bullet(
        "Đọc trạng thái từng gate từ data/acceptance-report.json; không suy diễn pass từ việc lệnh tồn tại.",
        "QA: ",
    )
    add_bullet(
        "Phân biệt kiểm thử tự động, technical review nội bộ và independent review.",
        "Thẩm quyền: ",
    )
    add_bullet(
        "Không thực hiện nghiên cứu thực nghiệm về hiệu quả học tập trong phạm vi báo cáo này.",
        "Ngoài phạm vi: ",
    )
    add_block_diagram(
        doc,
        [
            "DOCX nguồn chuẩn",
            "Trích xuất và chuẩn hóa",
            "Runtime offline",
            f"{gate_summary['total']} cổng QA",
            "Candidate",
            "Đánh giá độc lập",
        ],
        "Luồng tạo học liệu, kiểm soát chất lượng và chuyển giao nghiệm thu.",
    )

    # Chapter 1
    add_h1("Chương 1: Tổng quan hệ thống và kiến trúc kỹ thuật", new_page=True)
    add_h2("1.1. Kiến trúc và nguồn chuẩn")
    add_body(
        "Giáo trình vận hành bằng HTML, CSS và JavaScript tĩnh qua file://, USB hoặc static server. "
        "DOCX giữ vai trò nguồn narrative chuẩn; các fragment, bundle và manifest sinh tự động "
        "không được sửa tay."
    )
    add_bullet(
        "Runtime không yêu cầu backend hoặc CDN; npm chỉ phục vụ phát triển và QA.",
        "Ngoại tuyến: ",
    )
    add_bullet(
        "KaTeX, Three.js và PDF.js được đóng gói cục bộ trong candidate.",
        "Phụ thuộc: ",
    )
    add_bullet(
        "Dữ liệu người học như tiến độ, ghi chú và quiz attempt lưu cục bộ trong trình duyệt.",
        "Dữ liệu người dùng: ",
    )
    add_image_box(
        "images/ch1/hinh-002.png",
        "Mô hình liên kết lực và sơ đồ phân tích vật thể tự do trong Tĩnh học.",
        "Sơ đồ vật thể tự do và các lực liên kết dùng trong nội dung Chương 1.",
        12.0,
    )

    add_h2("1.2. Quy mô nội dung")
    chapter_counts = evidence["chapterCounts"]
    add_body(
        f"Hệ thống có {len(evidence['routes'])} route: Chương 1 có {chapter_counts[1]}, "
        f"Chương 2 có {chapter_counts[2]}, Chương 3 có {chapter_counts[3]} và "
        f"{evidence['otherRoutes']} route dẫn nhập/tra cứu."
    )
    route_rows = [
        ("Chương 1", "Tĩnh học", chapter_counts[1]),
        ("Chương 2", "Động học", chapter_counts[2]),
        ("Chương 3", "Động lực học", chapter_counts[3]),
        ("Bổ trợ", "Lời nói đầu, tác giả, tài liệu tham khảo", evidence["otherRoutes"]),
    ]
    add_caption(doc, "Bảng", "Table", "Phân bố route theo chương.")
    add_data_table(doc, ["Phần", "Phạm vi", "Số route"], route_rows, [Cm(3.0), Cm(10.0), Cm(3.0)])

    # Chapter 2
    add_h1("Chương 2: Quy cách trình bày và chuẩn tiêu chí", new_page=True)
    add_h2("2.1. Toán học ngữ nghĩa")
    add_body(
        f"Content manifest ghi {evidence['equationOccurrenceCount']} lần xuất hiện công thức; "
        f"registry ngữ nghĩa hiện có {evidence['equationCount']} hàng ánh xạ. "
        "Hai chỉ số phục vụ hai mục đích khác nhau và không được dùng thay thế nhau."
    )
    add_bullet(
        "KaTeX/MathML cung cấp hiển thị vector sắc nét và nội dung ngữ nghĩa trong phạm vi mapping.",
        "Hiển thị: ",
    )
    add_bullet(
        "Extractor loại placeholder '(.)' và strict tests bảo vệ hồi quy.",
        "Làm sạch: ",
    )
    add_bullet(
        "Technical PASS không thay thế kiểm tra ý nghĩa, đơn vị và ngữ cảnh của chuyên gia cơ học.",
        "Giới hạn: ",
    )
    add_image_box(
        "images/ch1/hinh-078.png",
        "Sơ đồ dầm chịu tải trọng phân bố và biểu đồ nội lực.",
        "Dầm chịu tải phân bố và biểu đồ nội lực dùng làm mẫu kiểm tra trình bày công thức–hình.",
        13.0,
    )

    add_h2("2.2. Ảnh động và phương án giảm chuyển động")
    add_body(
        f"{len(GIF_ROWS)} GIF phát hành được ánh xạ sang ảnh PNG canonical. Runtime chuyển "
        "về PNG khi người dùng bật prefers-reduced-motion hoặc khi GIF không tải được."
    )
    add_caption(
        doc,
        "Bảng",
        "Table",
        f"Danh mục {len(GIF_ROWS)} GIF và ảnh PNG dự phòng.",
    )
    gif_table_rows = [
        (index, gif_path, description, fallback)
        for index, (gif_path, description, fallback) in enumerate(GIF_ROWS, 1)
    ]
    add_data_table(
        doc,
        ["STT", "GIF", "Hiện tượng", "PNG dự phòng"],
        gif_table_rows,
        [Cm(1.0), Cm(4.6), Cm(6.5), Cm(4.3)],
        font_size=7.8,
    )
    add_image_box(
        "images/ch2/hinh-196.png",
        "Cơ cấu truyền động bánh răng ăn khớp trong.",
        "Hình cơ cấu truyền động bánh răng thuộc nội dung Chương 2.",
        12.5,
    )

    add_h2("2.3. Trắc nghiệm và gói LMS")
    quiz_total = sum(evidence["quizCounts"].values())
    quiz_distribution = ", ".join(
        f"Chương {chapter}: {count}"
        for chapter, count in sorted(evidence["quizCounts"].items())
    )
    lms_gate_status = gate_status["lms-adapters"]
    lms_result = (
        "Gate adapter QTI 3/Common Cartridge 1.4 đã pass kiểm tra cục bộ."
        if lms_gate_status == "pass"
        else f"Gate adapter LMS có trạng thái {lms_gate_status}; không tuyên bố đã đạt."
    )
    add_body(
        f"Kho câu hỏi có {quiz_total} mục ({quiz_distribution}). {lms_result}"
    )
    add_bullet(
        "Không có target LMS hoặc execution evidence; chưa tuyên bố nhập thành công vào "
        "Canvas, Moodle hay Blackboard.",
        "Giới hạn liên thông: ",
    )
    add_bullet(
        f"Trạng thái data/lms-targets.json: {lms_targets['status']}.",
        "Nguồn trạng thái: ",
    )

    add_h2("2.4. Khả năng tiếp cận")
    accessibility_gate_status = gate_status["phase-08-accessibility"]
    if accessibility_gate_status == "pass":
        accessibility_result = (
            f"Automation bằng {accessibility['automation']['runner']} trên "
            f"{accessibility['automation']['transport']} có trạng thái "
            f"{accessibility['automation']['status']}."
        )
    else:
        accessibility_result = (
            f"Gate accessibility automation có trạng thái "
            f"{accessibility_gate_status}; không tuyên bố đã đạt."
        )
    add_body(
        f"{accessibility_result} Manual review hiện là "
        f"{accessibility['manualReview']['status']}."
    )
    add_bullet(
        "Automation bao phủ shell, search, quiz, PDF chrome và một số điều khiển mô phỏng đại diện.",
        "Đã kiểm: ",
    )
    add_bullet(
        "Screen reader, text spacing, focus-obscured, motion comprehension và scientific "
        "visualization equivalence còn cần đánh giá thủ công.",
        "Chưa đủ bằng chứng: ",
    )
    add_bullet(
        "Báo cáo không tuyên bố tuân thủ WCAG 2.2 AA toàn hệ thống.",
        "Phạm vi tuyên bố: ",
    )
    add_image_box(
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-02-trang-chu-mobile-390x844.png",
        "Giao diện trang chủ tại viewport 390 × 844, bằng chứng responsive đại diện.",
        "Ảnh chụp giao diện mobile 390 × 844; bằng chứng trình bày, không phải chứng nhận WCAG.",
        6.5,
    )

    # Chapter 3
    add_h1("Chương 3: Phương pháp luận và kiến trúc mô phỏng", new_page=True)
    add_h2("3.1. Giới hạn khái niệm 4D")
    add_body(
        "Trong dự án, 4D là khung mô tả học tập, không phải chiều không gian thứ tư hoặc "
        "một runtime độc lập. Khung gồm ba thành phần đồng thời:"
    )
    add_bullet(
        "Biểu diễn 3D khi chiều sâu làm rõ quan hệ cơ học mà 2D khó phân biệt.",
        "Không gian: ",
    )
    add_bullet(
        "Tiến triển theo thời gian hoặc trạng thái qua tham số, bước mô phỏng và reset.",
        "Thời gian/trạng thái: ",
    )
    add_bullet(
        "Điều khiển, drag handle hoặc chuyển giữa Sim2 và Sim3.",
        "Tương tác: ",
    )

    add_h2("3.2. Sim2 canonical")
    simulation_gate_status = gate_status["simulation-evidence-currentness"]
    simulation_evidence_text = (
        "Gate currentness mô phỏng đã pass."
        if simulation_gate_status == "pass"
        else f"Gate currentness mô phỏng có trạng thái {simulation_gate_status}; "
        "các mô tả sau chỉ là inventory."
    )
    add_body(
        f"Registry Sim2 có {len(evidence['simulationSpecs'])} route SVG-first. "
        f"{simulation_evidence_text} Đồng hồ fixed-step 1/60 s cung cấp nhịp cập nhật ổn định; "
        "phương pháp tính vật lý phụ thuộc từng route. RK4 chỉ dùng ở bài toán ODE phù hợp, "
        "không phải mặc định cho mọi mô phỏng."
    )
    add_bullet("Biến đổi world-to-screen và responsive CSS scale.", "Hình học: ")
    add_bullet("Native controls, drag handle và readout cùng dùng một state.", "Tương tác: ")
    add_bullet("dispose() dọn listener, observer, RAF và DOM thuộc route.", "Vòng đời: ")
    add_image_box(
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch1-6-3-negative-area-win32.png",
        "Mô phỏng xác định trọng tâm hình phẳng ghép và khoét, route ch1-6-3.",
        "Ảnh chụp mô phỏng Sim2 route ch1-6-3 với hình ghép và phần diện tích khoét.",
    )

    add_h2("3.3. Sim3 pilot và fallback")
    sim3_gate_status = gate_status["sim3-pilot"]
    sim3_evidence_text = (
        "Gate Sim3 pilot đã pass."
        if sim3_gate_status == "pass"
        else f"Gate Sim3 pilot có trạng thái {sim3_gate_status}; không tuyên bố đã đạt."
    )
    add_body(
        f"Registry Sim3 có {len(evidence['sim3Reviews'])} adapter pilot. "
        f"{sim3_evidence_text} Đây là lớp tùy chọn; Sim2 vẫn là đường chạy canonical. "
        "Technical review không phải phê duyệt sư phạm độc lập."
    )
    add_bullet("Hệ tọa độ tay phải: +X phải, +Y lên, +Z hướng về người xem.", "Tọa độ: ")
    add_bullet("Render theo nhu cầu, cap DPR và giải phóng tài nguyên GPU.", "Hiệu năng: ")
    add_bullet("WebGL/setup/render lỗi phải thông báo tiếng Việt và trở về Sim2.", "Fallback: ")
    add_block_diagram(
        doc,
        [
            "Sim2 canonical",
            "Đánh giá giá trị chiều sâu",
            "Bật Sim3 pilot",
            "Theo dõi lỗi/nhận thức",
            "Fallback Sim2",
        ],
        "Quan hệ Sim2 canonical, Sim3 pilot và đường fallback.",
    )
    add_image_box(
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch2-4-4-coriolis-callout-win32.png",
        "Mô phỏng hợp chuyển động và gia tốc Coriolis, route ch2-4-4.",
        "Ảnh chụp route ch2-4-4 minh họa vectơ trong bài toán gia tốc Coriolis.",
    )

    # Chapter 4
    add_h1("Chương 4: Demo hệ thống và danh mục mô phỏng", new_page=True)
    add_h2("4.1. Kịch bản demo hệ thống trong 90 giây")
    add_body(
        "Kịch bản demo được giới hạn ở một chuỗi quan sát có thể kiểm chứng: mở gói ngoại tuyến, "
        "điều hướng đến bài Mô men, thao tác Sim2, đọc đầu ra và đối chiếu bản PDF cục bộ."
    )
    add_block_diagram(
        doc,
        [
            "Mở package qua file://",
            "Chương 1 › Mô men",
            "F = 50 N; d = 4,00 m",
            "M = 200 N·m",
            "Đối chiếu PDF",
        ],
        "Luồng demo 90 giây từ gói ngoại tuyến đến đối chiếu nội dung.",
    )
    demo_rows = [
        ("00:00–00:15", "Mở gói", "package/index.html qua file://"),
        ("00:15–00:30", "Vào bài", "Chương 1 › I › 4. Mô men"),
        ("00:30–01:00", "Thao tác", "Giữ F = 50 N; kéo d đến 4,00 m"),
        ("01:00–01:15", "Quan sát", "Readout M = 200 N·m; chiều quay cập nhật"),
        ("01:15–01:30", "Đối chiếu", "Mở PDF cục bộ và quay lại bài"),
    ]
    add_caption(doc, "Bảng", "Table", "Kịch bản demo hệ thống và tiêu chí quan sát.")
    add_data_table(
        doc,
        ["Thời gian", "Bước", "Quan sát bắt buộc"],
        demo_rows,
        [Cm(3.0), Cm(3.0), Cm(10.4)],
        font_size=8.5,
    )
    add_image_box(
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-01-trang-chu-desktop-1440x1000.png",
        "Bước 1: mở trang chủ candidate từ gói ngoại tuyến.",
        "Trang chủ giáo trình điện tử dùng làm bước mở đầu demo.",
        13.0,
    )
    add_image_box(
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-04-mo-men-ch1-1-4-1440x1000.png",
        "Bước 2–4: thao tác route ch1-1-4; F = 50 N, d = 4,00 m, M = 200 N·m.",
        "Mô phỏng mô men lực với đầu vào, mô hình và readout đầu ra.",
        13.0,
    )
    add_image_box(
        "assets/designs/bao-cao-nghiem-thu-giao-trinh-dien-tu/assets/img-06-pdf-viewer-1440x1000.png",
        "Bước 5: mở PDF cục bộ để đối chiếu nội dung nguồn.",
        "PDF viewer cục bộ dùng trong bước đối chiếu cuối demo.",
        13.0,
    )

    add_h2("4.2. Danh mục mô phỏng đã công bố")
    add_body(
        "Danh mục dưới đây sinh từ data/simulation-specifications.json. Cột tương tác chỉ mô tả "
        "bề mặt người học; tính đúng học thuật vẫn cần reviewer độc lập."
    )
    sim3_ids = {record["id"] for record in evidence["sim3Reviews"]}
    sim_rows = []
    for spec in evidence["simulationSpecs"]:
        engine = "Sim2 SVG"
        if spec["id"] in sim3_ids:
            engine += " + Sim3 pilot"
        sim_rows.append(
            (
                spec["id"],
                f"Chương {spec['chapter']}",
                spec["title"],
                engine,
                SIM_INTERACTIONS.get(spec["id"], "Điều khiển theo route"),
            )
        )
    add_caption(
        doc,
        "Bảng",
        "Table",
        f"Danh mục {len(evidence['simulationSpecs'])} route Sim2 và "
        f"{len(evidence['sim3Reviews'])} adapter Sim3 pilot.",
    )
    add_data_table(
        doc,
        ["Route", "Chương", "Bài toán", "Động cơ", "Tương tác"],
        sim_rows,
        [Cm(2.3), Cm(1.8), Cm(5.2), Cm(3.1), Cm(4.0)],
        font_size=7.8,
    )
    add_image_box(
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch2-3-2-transmission-win32.png",
        "Mô phỏng truyền động bánh răng, route ch2-3-2.",
        "Ảnh chụp mô phỏng cơ cấu truyền động tại route ch2-3-2.",
    )
    add_image_box(
        "tools/sim2-visual/selective-baseline.spec.js-snapshots/ch3-6-2-collision-after-win32.png",
        "Mô phỏng va chạm với hệ số phục hồi, route ch3-6-2.",
        "Ảnh chụp trạng thái sau va chạm tại route ch3-6-2.",
    )

    # Chapter 5
    add_h1("Chương 5: Kết quả định lượng và ma trận QA", new_page=True)
    add_h2("5.1. Candidate")
    add_body(
        f"Candidate {release['releaseVersion']} chứa {release['staging']['fileCount']} tệp; "
        f"gói ZIP có dung lượng {format_megabytes(release['package']['sizeBytes'])}."
    )
    add_bullet(release["package"]["sha256"], "SHA-256: ")
    add_bullet(
        f"{len(evidence['routes'])} route; {quiz_total} câu hỏi; {len(GIF_ROWS)} GIF.",
        "Inventory ràng buộc candidate: ",
    )
    add_bullet(
        f"{evidence['equationOccurrenceCount']} lần xuất hiện công thức trong content manifest; "
        f"{evidence['equationCount']} hàng mapping ngữ nghĩa; "
        f"{len(evidence['simulationSpecs'])} Sim2 và "
        f"{len(evidence['sim3Reviews'])} Sim3 pilot.",
        "Registry kỹ thuật có hash evidence: ",
    )

    add_h2(f"5.2. Snapshot trạng thái {gate_summary['total']} cổng")
    summary_rows = [
        ("Pass", gate_summary["pass"]),
        ("Fail", gate_summary["fail"]),
        ("Blocked", gate_summary["blocked"]),
        ("Not run", gate_summary["notRun"]),
        ("Tổng", gate_summary["total"]),
    ]
    add_caption(doc, "Bảng", "Table", "Tổng hợp trạng thái cổng QA trong snapshot.")
    add_data_table(doc, ["Trạng thái", "Số lượng"], summary_rows, [Cm(8.0), Cm(4.0)])
    add_block_diagram(
        doc,
        [
            f"{gate_summary['total']} cổng",
            f"{gate_summary['pass']} pass",
            f"{gate_summary['fail']} fail + {gate_summary['blocked']} blocked + "
            f"{gate_summary['notRun']} not run",
            f"Decision: {acceptance['releaseDecision']['decision']}",
            f"Report: {profile['key']}",
        ],
        "Luồng quyết định từ snapshot QA và kiểm tra ràng buộc candidate.",
    )

    gate_rows = [
        (
            gate["gateId"],
            gate["status"],
            gate["owner"],
            Path(gate["artifact"]).name,
        )
        for gate in acceptance["gates"]
    ]
    add_caption(
        doc,
        "Bảng",
        "Table",
        f"Chi tiết {gate_summary['total']} cổng và bằng chứng quan sát.",
    )
    add_data_table(
        doc,
        ["Gate ID", "Trạng thái", "Chủ sở hữu", "Artifact"],
        gate_rows,
        [Cm(4.2), Cm(2.0), Cm(4.5), Cm(5.7)],
        font_size=7.4,
    )

    # Chapter 6
    add_h1("Chương 6: Phản biện khoa học và hướng hoàn thiện", new_page=True)
    add_h2("6.1. Kết quả đã xác minh")
    add_bullet(
        f"Snapshot acceptance có {gate_summary['pass']} cổng pass; "
        f"trạng thái binding={profile['key']}.",
        "Bằng chứng: ",
    )
    add_bullet(
        (
            "Simulation currentness và Sim3 pilot đều pass trong snapshot."
            if simulation_gate_status == "pass" and sim3_gate_status == "pass"
            else f"Simulation currentness={simulation_gate_status}; "
            f"Sim3 pilot={sim3_gate_status}; không tuyên bố đã đạt."
        ),
        "Mô phỏng: ",
    )
    add_bullet(
        (
            "Adapter LMS pass kiểm tra cục bộ trong snapshot."
            if lms_gate_status == "pass"
            else f"Adapter LMS={lms_gate_status}; không tuyên bố đã đạt."
        ),
        "Liên thông: ",
    )

    unresolved = [
        gate for gate in acceptance["gates"]
        if gate["status"] in {"fail", "blocked", "not-run", "notRun"}
    ]
    unresolved_count = len(unresolved) + (0 if binding["current"] else 1)
    add_h2(f"6.2. {unresolved_count} điều kiện hoặc sai lệch còn mở")
    if not binding["current"]:
        add_bullet(
            "Acceptance snapshot không ràng buộc candidate hiện hành; "
            + "; ".join(binding_issues),
            "candidate-evidence-binding: ",
        )
    for gate in unresolved:
        add_bullet(
            f"{gate['status']}; owner={gate['owner']}; artifact={gate['artifact']}.",
            f"{gate['gateId']}: ",
        )

    add_h2("6.3. Giới hạn khoa học")
    add_bullet(
        "Không có quyết định independent SME trong data/academic_signoffs.json.",
        "Học thuật: ",
    )
    add_bullet(
        "Learning outcomes và legal register còn provisional.",
        "Quản trị: ",
    )
    add_bullet(
        "Không có nghiên cứu đối chứng hoặc dữ liệu trước–sau để kết luận hiệu quả học tập.",
        "Hiệu quả sư phạm: ",
    )
    add_bullet(
        "Không có ma trận CDIO/ABET được cơ sở đào tạo phê duyệt.",
        "Kiểm định: ",
    )
    add_bullet(
        "Không có target LMS hoặc bằng chứng import.",
        "Liên thông thực tế: ",
    )

    add_h2("6.4. Lộ trình ưu tiên")
    add_body(
        "Thứ tự xử lý phải đóng bằng chứng trước khi mở rộng tính năng:"
    )
    add_bullet(
        "Hoàn tất academic signoff, accessibility review, independent smoke và Word round-trip.",
        "P0 — Nghiệm thu: ",
    )
    add_bullet(
        "Chạy lại toàn bộ ma trận, rebuild acceptance report và tái sinh tài liệu từ generator.",
        "P1 — Đồng bộ: ",
    )
    add_bullet(
        "Thử import trên LMS được chỉ định nếu cần tuyên bố liên thông thực tế.",
        "P2 — LMS: ",
    )
    add_bullet(
        "Thiết kế nghiên cứu hiệu quả học tập trước mọi tuyên bố sư phạm.",
        "P3 — Đánh giá giáo dục: ",
    )
    add_bullet(
        "Không mở rộng thêm Sim3 nếu chưa chứng minh giá trị 3D vượt 2D và kiểm soát tải nhận thức.",
        "P4 — Mô phỏng: ",
    )

    add_h2("6.5. Kết luận")
    add_body(profile["conclusion"])

    # References and status registers
    add_h1("Tài liệu tham chiếu và nguồn bằng chứng", new_page=True)
    source_rows = [
        ("Trạng thái nghiệm thu", "data/acceptance-report.json", acceptance["overallStatus"]),
        ("Candidate", release_candidate["summaryPath"], release_candidate["status"]),
        (
            "Ràng buộc candidate–acceptance",
            "data/evidence-registry.json",
            profile["key"],
        ),
        ("Khả năng tiếp cận", "data/accessibility-baseline.json", accessibility["manualReview"]["status"]),
        ("LMS", "data/lms-targets.json", lms_targets["status"]),
        ("Học thuật", "docs/academic-certification.md", "provisional"),
        ("Mô phỏng 4D", "docs/simulation-4d.md", "technical-review verified"),
        ("Ma trận QA", "docs/qa-gate-matrix.md", "canonical definitions"),
        ("Chuẩn đầu ra", "data/learning-outcomes.json", evidence["learningOutcomes"]["status"]),
    ]
    for record in evidence["legal"]["records"]:
        source_rows.append(
            (
                record["title"],
                record["officialSource"],
                record["reviewStatus"],
            )
        )
    add_caption(doc, "Bảng", "Table", "Nguồn bằng chứng và trạng thái thẩm quyền.")
    add_data_table(
        doc,
        ["Chủ đề", "Nguồn", "Trạng thái"],
        source_rows,
        [Cm(4.0), Cm(8.5), Cm(3.9)],
        font_size=8.0,
    )

    # Signature slots are neutral until an authoritative institution supplies metadata.
    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = signature_table.cell(0, 0)
    right = signature_table.cell(0, 1)
    for cell in (left, right):
        cell.width = Cm(8.0)
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
    left_paragraph = left.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = left_paragraph.add_run(
        "NHÓM BIÊN SOẠN VÀ PHÁT TRIỂN\nKỸ THUẬT HỌC LIỆU SỐ\n\n\n\n"
        "(Ký và ghi rõ họ tên)"
    )
    set_run_font(left_run, size=10, color=COLOR_NAVY, bold=True)
    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_run = right_paragraph.add_run(
        "ĐƠN VỊ THẨM ĐỊNH ĐƯỢC CHỈ ĐỊNH\nTHEO QUYẾT ĐỊNH CHÍNH THỨC\n\n\n\n"
        "(Chỉ ký sau khi đủ bằng chứng bắt buộc)"
    )
    set_run_font(right_run, size=10, color=COLOR_RED, bold=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    normalize_docx_package(output_path)
    print(
        f"Report generated: {output_path} | snapshot={acceptance['overallStatus']} | "
        f"binding={profile['key']} | gates={gate_summary['pass']}/{gate_summary['total']}"
    )
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Generate the evidence-calibrated scientific report DOCX."
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    arguments = parser.parse_args()
    build_report(arguments.output)


if __name__ == "__main__":
    main()
