"""
Analyze DOCX structure for the static textbook pipeline.

Usage:
  python tools/analyze_docx.py --input CoHocLyThuyet_Full_New.docx
"""
import argparse
import collections
import os
import re
import sys
import zipfile

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")


CHAPTER_TITLES = {"TĨNH HỌC", "ĐỘNG HỌC", "ĐỘNG LỰC HỌC"}


def resolve_path(path):
    if os.path.isabs(path):
        return path
    return os.path.abspath(os.path.join(os.getcwd(), path))


def paragraph_style(para):
    return para.style.name if para.style else ""


def is_toc_style(style):
    return style.lower().startswith("toc")


def clean_text(text):
    return " ".join((text or "").strip().split())


def collect_outline(doc):
    outline = []
    in_content = False
    for idx, para in enumerate(doc.paragraphs):
        style = paragraph_style(para)
        text = clean_text(para.text)
        if not text or is_toc_style(style):
            continue

        upper = text.upper()
        if "LỜI NÓI ĐẦU" in upper:
            outline.append((idx, "special", "LỜI NÓI ĐẦU"))
            in_content = True
            continue
        if "TÀI LIỆU THAM KHẢO" in upper:
            outline.append((idx, "special", "TÀI LIỆU THAM KHẢO"))
            continue

        if style.startswith("Heading "):
            try:
                level = int(style.split()[-1])
            except ValueError:
                continue
            if level == 1 and upper not in CHAPTER_TITLES:
                continue
            if in_content:
                outline.append((idx, f"H{level}", text))
    return outline


def validate_outline(outline):
    h1 = [row for row in outline if row[1] == "H1"]
    h2 = [row for row in outline if row[1] == "H2"]
    h3 = [row for row in outline if row[1] == "H3"]
    missing = []
    for chapter in CHAPTER_TITLES:
        if not any(row[2].upper() == chapter for row in h1):
            missing.append(chapter)
    if missing:
        raise SystemExit(f"Missing chapter heading(s): {', '.join(missing)}")
    if len(h2) < 15:
        raise SystemExit(f"Too few H2 sections detected: {len(h2)}")
    if len(h3) < 50:
        raise SystemExit(f"Too few H3 subsections detected: {len(h3)}")


def media_stats(docx_path):
    with zipfile.ZipFile(docx_path, "r") as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        ext_counts = collections.Counter(os.path.splitext(n)[1].lower() or "(none)" for n in media)
        doc_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    ole_progids = collections.Counter(re.findall(r'ProgID="([^"]+)"', doc_xml))
    return {
        "media": len(media),
        "extensions": ext_counts,
        "omml": doc_xml.count("<m:oMath"),
        "drawing_blips": doc_xml.count("<a:blip"),
        "vml_images": doc_xml.count("<v:imagedata"),
        "ole_progids": ole_progids,
        "tables": doc_xml.count("<w:tbl"),
    }


def print_route_preview(outline):
    current_chapter = 0
    current_section = 0
    subsection = 0
    roman = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

    print("\n=== ROUTE PREVIEW ===")
    for _, level, text in outline:
        if level == "H1":
            current_chapter += 1
            current_section = 0
            print(f"ch{current_chapter:<7} | {text}")
        elif level == "H2" and current_chapter:
            current_section += 1
            subsection = 0
            r = roman[current_section] if current_section < len(roman) else str(current_section)
            print(f"ch{current_chapter}-{current_section:<5} | {r}. {text}")
        elif level == "H3" and current_chapter and current_section:
            subsection += 1
            print(f"ch{current_chapter}-{current_section}-{subsection:<3} | {subsection}. {text}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="CoHocLyThuyet_Full_New.docx")
    parser.add_argument("--routes", action="store_true", help="Print generated route preview")
    args = parser.parse_args()

    docx_path = resolve_path(args.input)
    if not os.path.exists(docx_path):
        raise SystemExit(f"Input DOCX not found: {docx_path}")

    doc = Document(docx_path)
    outline = collect_outline(doc)
    validate_outline(outline)
    stats = media_stats(docx_path)
    style_counts = collections.Counter(paragraph_style(p) or "None" for p in doc.paragraphs)

    print("=== DOCX ANALYSIS ===")
    print(f"Input: {docx_path}")
    print(f"Size MB: {os.path.getsize(docx_path) / 1024 / 1024:.2f}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Sections: {len(doc.sections)}")
    print(f"Media files: {stats['media']}")
    print(f"OMML math objects: {stats['omml']}")
    print(f"Drawing images: {stats['drawing_blips']}")
    print(f"VML images: {stats['vml_images']}")
    print(f"OLE objects: {sum(stats['ole_progids'].values())}")
    print(f"XML tables: {stats['tables']}")

    if stats["ole_progids"]:
        print("\n=== OLE PROGIDS ===")
        for prog_id, count in stats["ole_progids"].most_common():
            print(f"{prog_id}: {count}")

    print("\n=== MEDIA EXTENSIONS ===")
    for ext, count in sorted(stats["extensions"].items()):
        print(f"{ext}: {count}")

    print("\n=== TOP PARAGRAPH STYLES ===")
    for style, count in style_counts.most_common(20):
        print(f"{count:5d} | {style}")

    print("\n=== CONTENT OUTLINE ===")
    for idx, level, text in outline:
        print(f"[{idx:04d}] {level:8s} | {text}")

    if args.routes:
        print_route_preview(outline)


if __name__ == "__main__":
    main()
