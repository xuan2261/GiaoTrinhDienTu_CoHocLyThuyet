"""
Extract DOCX content into static HTML fragments and chapter images.

Source of truth: CoHocLyThuyet_Full_New.docx

Usage:
  python tools/extract_docx.py --input CoHocLyThuyet_Full_New.docx --write
  python tools/extract_docx.py --input CoHocLyThuyet_Full_New.docx
"""
import argparse
import hashlib
import html
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile

from docx import Document
from lxml import etree

sys.stdout.reconfigure(encoding="utf-8")


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}

CHAPTER_NAMES = {1: "Tĩnh học", 2: "Động học", 3: "Động lực học"}
CHAPTER_HEADINGS = {"TĨNH HỌC": 1, "ĐỘNG HỌC": 2, "ĐỘNG LỰC HỌC": 3}
ROMAN = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
OMML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\OMML2MML.XSL"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"


def resolve_path(path, base=None):
    if os.path.isabs(path):
        return path
    return os.path.abspath(os.path.join(base or os.getcwd(), path))


def clean_text(text):
    return " ".join((text or "").strip().split())


def normalize_mathml(value):
    mathml = str(value or "").strip()
    if not mathml:
        return ""
    mathml = re.sub(r"<\?xml[^>]*\?>", "", mathml).strip()
    start = mathml.find("<math")
    close = "</math>"
    end = mathml.rfind(close)
    if start == -1:
        start = mathml.find("<mml:math")
        close = "</mml:math>"
        end = mathml.rfind(close)
    if start == -1 or end == -1:
        return ""
    mathml = mathml[start:end + len(close)].strip()
    try:
        root = etree.fromstring(mathml.encode("utf-8"))
    except etree.XMLSyntaxError:
        return ""
    if not etree.QName(root).localname == "math":
        return ""
    for element in root.iter():
        if not isinstance(element.tag, str):
            continue
        qname = etree.QName(element)
        if qname.namespace == MATHML_NS:
            element.tag = qname.localname
    etree.cleanup_namespaces(root)
    root.set("xmlns", MATHML_NS)
    return etree.tostring(root, encoding="unicode", pretty_print=False)


def mathml_with_display(mathml, display):
    try:
        root = etree.fromstring(str(mathml or "").encode("utf-8"))
    except etree.XMLSyntaxError:
        return mathml
    if display:
        root.set("display", "block")
    else:
        root.attrib.pop("display", None)
    return etree.tostring(root, encoding="unicode", pretty_print=False)


def para_style(para):
    return para.style.name if para.style else ""


def heading_level(para):
    style = para_style(para)
    if style.startswith("Heading "):
        try:
            return int(style.split()[-1])
        except ValueError:
            return 0
    return 0


def is_toc(para):
    return para_style(para).lower().startswith("toc")


def package_data(docx_path):
    with zipfile.ZipFile(docx_path, "r") as zf:
        doc_xml = zf.read("word/document.xml")
        rels_xml = zf.read("word/_rels/document.xml.rels")
        media_blobs = {
            os.path.basename(name): zf.read(name)
            for name in zf.namelist()
            if name.startswith("word/media/")
        }

    rels = etree.fromstring(rels_xml)
    rid_to_media = {}
    for rel in rels:
        rid = rel.get("Id")
        target = rel.get("Target", "")
        if rid and "media/" in target:
            rid_to_media[rid] = os.path.basename(target)

    root = etree.fromstring(doc_xml)
    body = root.find("w:body", NS)
    xml_paragraphs = [node for node in body if node.tag == f"{{{NS['w']}}}p"]
    return xml_paragraphs, rid_to_media, media_blobs


def qn(prefix, name):
    return f"{{{NS[prefix]}}}{name}"


def node_name(node):
    return etree.QName(node).localname if isinstance(node.tag, str) else ""


def parse_style_dimensions(style):
    dims = {}
    for key in ("width", "height"):
        match = re.search(rf"{key}\s*:\s*([0-9.]+)(pt|in|cm|mm|px)", style or "", flags=re.I)
        if match:
            dims[key] = (float(match.group(1)), match.group(2).lower())
    return dims


def dimension_to_pt(value):
    if not value:
        return None
    amount, unit = value
    if unit == "pt":
        return amount
    if unit == "in":
        return amount * 72
    if unit == "cm":
        return amount * 28.3465
    if unit == "mm":
        return amount * 2.83465
    if unit == "px":
        return amount * 0.75
    return None


def closest(node, tag):
    current = node
    while current is not None:
        if current.tag == tag:
            return current
        current = current.getparent()
    return None


def first_descendant(node, xpath):
    found = node.xpath(xpath, namespaces=NS)
    return found[0] if found else None


def image_dimensions(image_node):
    shape = closest(image_node, qn("v", "shape"))
    if shape is not None:
        dims = parse_style_dimensions(shape.get("style", ""))
        return dimension_to_pt(dims.get("width")), dimension_to_pt(dims.get("height"))

    inline = closest(image_node, qn("wp", "inline"))
    if inline is None:
        inline = closest(image_node, qn("wp", "anchor"))
    if inline is not None:
        extent = first_descendant(inline, ".//wp:extent")
        if extent is not None:
            try:
                return int(extent.get("cx", "0")) / 12700, int(extent.get("cy", "0")) / 12700
            except ValueError:
                return None, None
    return None, None


def ole_prog_id(image_node):
    current = image_node
    while current is not None:
        ole = first_descendant(current, ".//o:OLEObject")
        if ole is not None:
            return ole.get("ProgID")
        if current.tag == qn("w", "p"):
            break
        current = current.getparent()
    return None


def media_from_image_node(image_node, rid_to_media):
    rid = None
    if image_node.tag == qn("a", "blip"):
        rid = image_node.get(qn("r", "embed"))
    elif image_node.tag == qn("v", "imagedata"):
        rid = image_node.get(qn("r", "id"))
    return rid, rid_to_media.get(rid)


def equation_number_only(text):
    normalized = clean_text(text)
    return bool(re.fullmatch(r"[\(\[]?\d+(?:\.\d+)*[\)\]]?|\(\.\)", normalized))


def classify_media(prog_id, width_pt, height_pt, paragraph_text, media_name):
    if not media_name:
        return "unknown"
    if prog_id and prog_id.startswith("Equation."):
        if not clean_text(paragraph_text) or equation_number_only(paragraph_text):
            return "math-display"
        if (width_pt and width_pt >= 180) or (height_pt and height_pt >= 80):
            return "math-display"
        return "math-inline"
    return "figure"


def load_omml_transformer():
    if not os.path.exists(OMML_XSL):
        return None
    try:
        return etree.XSLT(etree.parse(OMML_XSL))
    except Exception:
        return None


def paragraph_math(xml_para, transformer):
    blocks = []
    for node in xml_para.findall(".//m:oMath", NS):
        if transformer is None:
            blocks.append('<div class="mathml-block"><em>[Công thức OMML]</em></div>')
            continue
        try:
            math_node = etree.fromstring(etree.tostring(node))
            result = transformer(math_node)
            mathml = normalize_mathml(result)
            if mathml:
                blocks.append(f'<div class="mathml-block">{mathml}</div>')
        except Exception as exc:
            blocks.append(
                '<div class="mathml-block math-fallback">'
                f'<em>[Không chuyển được công thức: {html.escape(str(exc))}]</em>'
                "</div>"
            )
    return blocks


def count_omml(xml_paragraphs):
    return sum(len(xml_para.findall(".//m:oMath", NS)) for xml_para in xml_paragraphs)


def conversion_required(media_blobs):
    browser_exts = {".png", ".jpg", ".jpeg", ".gif"}
    return [
        name for name in media_blobs
        if os.path.splitext(name)[1].lower() not in browser_exts
    ]


def preflight_dependencies(media_blobs, omml_count, transformer, write):
    convert_media = conversion_required(media_blobs)
    magick = shutil.which("magick")
    if write and convert_media and not magick:
        raise SystemExit(
            "ImageMagick 'magick' is required before writing because "
            f"{len(convert_media)} DOCX media file(s) need conversion."
        )
    if write and omml_count and transformer is None:
        raise SystemExit(
            "OMML2MML.XSL is required before writing because "
            f"{omml_count} OMML math object(s) were found."
        )
    if convert_media and not magick:
        print(f"[WARN] {len(convert_media)} media file(s) need ImageMagick conversion.")
    if omml_count and transformer is None:
        print(f"[WARN] {omml_count} OMML math object(s) need OMML2MML.XSL conversion.")


def load_equation_mapping(root):
    path = os.path.join(root, "data", "equation_mapping.json")
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as fh:
            raw = json.load(fh)
    except (OSError, json.JSONDecodeError):
        return {}
    items = raw.get("equations", raw) if isinstance(raw, dict) else raw
    mapping = {}
    for item in items if isinstance(items, list) else []:
        if not item.get("reviewed"):
            continue
        media_hash = item.get("hash")
        if media_hash and (item.get("latex") or item.get("mathml") or item.get("artifact")):
            normalized = dict(item)
            if normalized.get("mathml"):
                mathml = normalize_mathml(normalized.get("mathml"))
                if not mathml:
                    continue
                normalized["mathml"] = mathml
            mapping[media_hash] = normalized
    return mapping


def render_mapped_equation(mapping, fallback_kind, rel=None, alt="", style=""):
    display = fallback_kind == "math-display"
    artifact = mapping.get("artifact")
    if artifact == "blank":
        return {"kind": "block" if display else "inline", "html": ""}
    if artifact == "figure" and rel:
        safe_alt = html.escape(mapping.get("alt") or alt or "Hình minh họa")
        return {
            "kind": "block",
            "html": f'<div class="figure-container"><img src="{rel}" alt="{safe_alt}" loading="lazy"{style}></div>',
        }
    if mapping.get("mathml"):
        wrapper = "div" if display else "span"
        cls = "mathml-block" if display else "mathml-inline"
        mathml = mathml_with_display(mapping["mathml"], display)
        return {"kind": "block" if display else "inline", "html": f'<{wrapper} class="{cls}">{mathml}</{wrapper}>'}
    latex = html.escape(mapping.get("latex", ""))
    if display:
        return {"kind": "block", "html": f'<div class="math-tex-block">\\[{latex}\\]</div>'}
    return {"kind": "inline", "html": f'<span class="math-tex">\\({latex}\\)</span>'}


class ImageWriter:
    def __init__(self, root, media_blobs, write, equation_mapping=None):
        self.root = root
        self.media_blobs = media_blobs
        self.write = write
        self.magick = shutil.which("magick")
        self.counters = {0: 0, 1: 0, 2: 0, 3: 0}
        self.saved = {0: {}, 1: {}, 2: {}, 3: {}}
        self.failures = []
        self.report = []
        self.equation_report = []
        self.class_counts = {"math-inline": 0, "math-display": 0, "figure": 0, "unknown": 0}
        self.equation_mapping = equation_mapping or {}

    def prepare(self):
        if not self.write:
            return
        for chapter in (0, 1, 2, 3):
            chapter_dir = os.path.join(self.root, "images", f"ch{chapter}")
            if os.path.exists(chapter_dir):
                shutil.rmtree(chapter_dir)
            os.makedirs(chapter_dir, exist_ok=True)

    def media_hash(self, media_name):
        blob = self.media_blobs.get(media_name)
        if blob is None:
            return None
        return hashlib.sha256(blob).hexdigest()

    def asset_for(self, chapter, media_name):
        if chapter not in self.saved:
            return None
        if media_name in self.saved[chapter]:
            return self.saved[chapter][media_name]

        blob = self.media_blobs.get(media_name)
        if blob is None:
            self.failures.append({"chapter": chapter, "media": media_name, "reason": "missing blob"})
            return None

        ext = os.path.splitext(media_name)[1].lower()
        browser_ext = ".jpg" if ext in (".jpg", ".jpeg") else ".png"
        needs_convert = ext not in (".png", ".jpg", ".jpeg", ".gif")
        self.counters[chapter] += 1
        out_name = f"hinh-{self.counters[chapter]:03d}{browser_ext}"
        rel_path = f"images/ch{chapter}/{out_name}"

        if not self.write:
            self.saved[chapter][media_name] = rel_path
            return rel_path

        out_path = os.path.join(self.root, rel_path)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        try:
            if not needs_convert:
                with open(out_path, "wb") as fh:
                    fh.write(blob)
            else:
                if not self.magick:
                    raise RuntimeError("ImageMagick magick command not found")
                with tempfile.TemporaryDirectory() as tmp:
                    input_path = os.path.join(tmp, media_name)
                    converted_path = os.path.join(tmp, out_name)
                    with open(input_path, "wb") as fh:
                        fh.write(blob)
                    proc = subprocess.run(
                        [self.magick, input_path, converted_path],
                        text=True,
                        capture_output=True,
                        timeout=30,
                    )
                    if proc.returncode != 0 or not os.path.exists(converted_path):
                        msg = (proc.stderr or proc.stdout or "conversion failed").strip()
                        raise RuntimeError(msg)
                    shutil.copyfile(converted_path, out_path)
            self.saved[chapter][media_name] = rel_path
            self.report.append({"chapter": chapter, "media": media_name, "output": rel_path})
            return rel_path
        except Exception as exc:
            raw_dir = os.path.join(self.root, "images", f"ch{chapter}", "raw")
            os.makedirs(raw_dir, exist_ok=True)
            raw_path = os.path.join(raw_dir, media_name)
            with open(raw_path, "wb") as fh:
                fh.write(blob)
            self.failures.append({"chapter": chapter, "media": media_name, "reason": str(exc), "raw": raw_path})
            return None

    def render_image_segment(self, chapter, media_name, meta):
        kind = meta["kind"]
        self.class_counts[kind] = self.class_counts.get(kind, 0) + 1
        media_hash = self.media_hash(media_name) if media_name else None
        mapping = self.equation_mapping.get(media_hash) if media_hash else None
        rel = self.asset_for(chapter, media_name) if media_name else None

        record = {
            "chapter": chapter,
            "paragraph_index": meta["paragraph_index"],
            "media": media_name,
            "prog_id": meta["prog_id"],
            "kind": kind,
            "width_pt": round(meta["width_pt"], 2) if meta["width_pt"] else None,
            "height_pt": round(meta["height_pt"], 2) if meta["height_pt"] else None,
            "text_context": meta["text_context"],
            "output": rel,
            "hash": media_hash,
        }
        self.equation_report.append(record)

        if not rel:
            safe = html.escape(media_name or "unknown")
            if kind.startswith("math"):
                return {"kind": "inline", "html": f'<span class="math-missing">[Không chuyển được công thức: {safe}]</span>'}
            return {"kind": "block", "html": f'<div class="figure-missing">Không chuyển được hình: {safe}</div>'}

        width = meta["width_pt"]
        height = meta["height_pt"]
        style = ""
        if width and height and kind.startswith("math"):
            style = f' style="width:{width:.2f}pt;height:{height:.2f}pt"'

        if mapping:
            alt = html.escape(f"Công thức {media_name}")
            return render_mapped_equation(mapping, kind, rel=rel, alt=alt, style=style)

        if kind == "math-inline":
            alt = html.escape(f"Công thức {media_name}")
            return {
                "kind": "inline",
                "html": f'<img class="math-img math-img-inline" src="{rel}" alt="{alt}" loading="lazy"{style}>',
            }
        if kind == "math-display":
            alt = html.escape(f"Công thức {media_name}")
            return {
                "kind": "block",
                "html": (
                    f'<div class="math-img-block">'
                    f'<img class="math-img" src="{rel}" alt="{alt}" loading="lazy"{style}>'
                    "</div>"
                ),
            }

        alt = html.escape(f"Hình minh họa chương {chapter}")
        return {"kind": "block", "html": f'<div class="figure-container"><img src="{rel}" alt="{alt}" loading="lazy"></div>'}

    def write_report(self):
        if not self.write:
            return
        report_path = os.path.join(self.root, "tools", "image_mapping.json")
        data = {"saved": self.report, "failures": self.failures}
        with open(report_path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, indent=2)

    def write_equation_report(self, write_report):
        if not write_report:
            return
        report_path = os.path.join(self.root, "tools", "equation_report.json")
        data = {"counts": self.class_counts, "items": self.equation_report}
        with open(report_path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, indent=2)


def run_text(run):
    text = run.text
    if not text:
        return ""
    text = html.escape(text)
    if run.bold and run.italic:
        text = f"<strong><em>{text}</em></strong>"
    elif run.bold:
        text = f"<strong>{text}</strong>"
    elif run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    if getattr(run.font, "superscript", None):
        text = f"<sup>{text}</sup>"
    elif getattr(run.font, "subscript", None):
        text = f"<sub>{text}</sub>"
    return text


def paragraph_text_html(para):
    return "".join(run_text(run) for run in para.runs).strip()


def rpr_flag(rpr, name):
    if rpr is None:
        return False
    node = rpr.find(f"w:{name}", NS)
    if node is None:
        return False
    value = node.get(qn("w", "val"))
    return value not in ("0", "false", "False", "off", "none")


def format_run_xml_text(text, rpr):
    if not text:
        return ""
    if not text.strip():
        return html.escape(text)
    text = html.escape(text)
    bold = rpr_flag(rpr, "b")
    italic = rpr_flag(rpr, "i")
    underline = rpr_flag(rpr, "u")
    vert = None
    if rpr is not None:
        vert_node = rpr.find("w:vertAlign", NS)
        if vert_node is not None:
            vert = vert_node.get(qn("w", "val"))

    if bold and italic:
        text = f"<strong><em>{text}</em></strong>"
    elif bold:
        text = f"<strong>{text}</strong>"
    elif italic:
        text = f"<em>{text}</em>"
    if underline:
        text = f"<u>{text}</u>"
    if vert == "superscript":
        text = f"<sup>{text}</sup>"
    elif vert == "subscript":
        text = f"<sub>{text}</sub>"
    return text


def render_omml_segment(node, transformer, display=False):
    if transformer is None:
        tag = "div" if display else "span"
        cls = "mathml-block" if display else "mathml-inline"
        return {"kind": "block" if display else "inline", "html": f'<{tag} class="{cls}"><em>[Công thức OMML]</em></{tag}>'}
    try:
        math_node = etree.fromstring(etree.tostring(node))
        result = transformer(math_node)
        mathml = normalize_mathml(result)
        if not mathml:
            return None
        tag = "div" if display else "span"
        cls = "mathml-block" if display else "mathml-inline"
        return {"kind": "block" if display else "inline", "html": f'<{tag} class="{cls}">{mathml}</{tag}>'}
    except Exception as exc:
        safe = html.escape(str(exc))
        return {
            "kind": "block" if display else "inline",
            "html": f'<span class="math-fallback"><em>[Không chuyển được công thức: {safe}]</em></span>',
        }


def render_omml_segments(node, transformer):
    if node.tag == qn("m", "oMathPara"):
        segments = []
        for math_node in node.findall(".//m:oMath", NS):
            segment = render_omml_segment(math_node, transformer, display=True)
            if segment:
                segments.append(segment)
        return segments
    segment = render_omml_segment(node, transformer, display=False)
    return [segment] if segment else []


def render_image_node(image_node, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media):
    _rid, media_name = media_from_image_node(image_node, rid_to_media)
    width_pt, height_pt = image_dimensions(image_node)
    prog_id = ole_prog_id(image_node)
    kind = classify_media(prog_id, width_pt, height_pt, paragraph_text, media_name)
    meta = {
        "paragraph_index": paragraph_index,
        "prog_id": prog_id,
        "kind": kind,
        "width_pt": width_pt,
        "height_pt": height_pt,
        "text_context": clean_text(paragraph_text)[:180],
    }
    return image_writer.render_image_segment(chapter, media_name, meta)


def render_run_segments(run, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media, transformer):
    rpr = run.find("w:rPr", NS)
    segments = []
    text_buffer = []

    def flush_text():
        if text_buffer:
            text = "".join(text_buffer)
            html_text = format_run_xml_text(text, rpr)
            if html_text:
                segments.append({"kind": "inline", "html": html_text})
            text_buffer.clear()

    for child in run:
        if child.tag == qn("w", "rPr"):
            continue
        if child.tag == qn("w", "t"):
            text_buffer.append(child.text or "")
        elif child.tag == qn("w", "tab"):
            text_buffer.append(" ")
        elif child.tag == qn("w", "br"):
            flush_text()
            segments.append({"kind": "inline", "html": "<br>"})
        elif child.tag == qn("w", "sym"):
            char = child.get(qn("w", "char"))
            if char:
                try:
                    text_buffer.append(chr(int(char, 16)))
                except ValueError:
                    pass
        elif child.findall(".//a:blip", NS) or child.findall(".//v:imagedata", NS):
            flush_text()
            for image_node in child.findall(".//a:blip", NS):
                segments.append(render_image_node(image_node, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media))
            for image_node in child.findall(".//v:imagedata", NS):
                segments.append(render_image_node(image_node, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media))
        elif child.tag in (qn("m", "oMath"), qn("m", "oMathPara")):
            flush_text()
            segments.extend(render_omml_segments(child, transformer))

    flush_text()
    return segments


def paragraph_segments(xml_para, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media, transformer):
    segments = []
    for child in xml_para:
        if child.tag == qn("w", "pPr"):
            continue
        if child.tag == qn("w", "r"):
            segments.extend(render_run_segments(child, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media, transformer))
        elif child.tag == qn("w", "hyperlink"):
            for run in child.findall(".//w:r", NS):
                segments.extend(render_run_segments(run, chapter, paragraph_index, paragraph_text, image_writer, rid_to_media, transformer))
        elif child.tag in (qn("m", "oMath"), qn("m", "oMathPara")):
            segments.extend(render_omml_segments(child, transformer))
    return [segment for segment in segments if segment and segment.get("html")]


INLINE_MATH_CLASS_RE = re.compile(r'class="(?:mathml-inline|math-tex)\b')
TAG_RE = re.compile(r"<[^>]+>")


def visible_text(value):
    return html.unescape(TAG_RE.sub("", value or ""))


def starts_with_letter(value):
    return bool(re.search(r"^\s*[^\W\d_]", visible_text(value), flags=re.UNICODE))


def ends_with_letter(value):
    return bool(re.search(r"[^\W\d_]\s*$", visible_text(value), flags=re.UNICODE))


def is_inline_math(segment):
    return bool(INLINE_MATH_CLASS_RE.search(segment.get("html", "")))


def needs_inline_spacing(left, right):
    left_math = is_inline_math(left)
    right_math = is_inline_math(right)
    if left_math == right_math:
        return False
    if left_math:
        return starts_with_letter(right.get("html", ""))
    return ends_with_letter(left.get("html", ""))


def join_inline_segments(segments):
    parts = []
    previous = None
    for segment in segments:
        if previous and needs_inline_spacing(previous, segment):
            parts.append(" ")
        parts.append(segment["html"])
        previous = segment
    return "".join(parts).strip()


def inline_html(segments):
    return join_inline_segments([segment for segment in segments if segment["kind"] == "inline"])


def is_list_para(para):
    style = para_style(para)
    text = clean_text(para.text)
    return "List" in style or bool(re.match(r"^[-+•]\s+", text))


def strip_list_marker(text):
    return re.sub(r"^[-+•]\s+", "", text).strip()


def strip_list_marker_from_segments(segments):
    for segment in segments:
        if segment["kind"] == "inline":
            segment["html"] = strip_list_marker(segment["html"])
            return


def render_segments_as_blocks(segments, css=""):
    parts = []
    buffer = []
    for segment in segments:
        if segment["kind"] == "inline":
            buffer.append(segment)
            continue
        content = join_inline_segments(buffer)
        if content:
            parts.append(f"<p{css}>{content}</p>")
            buffer.clear()
        parts.append(segment["html"])
    content = join_inline_segments(buffer)
    if content:
        parts.append(f"<p{css}>{content}</p>")
    return parts


def render_segments_as_list_item(segments):
    item = ["  <li>"]
    buffer = []
    for segment in segments:
        if segment["kind"] == "inline":
            buffer.append(segment)
            continue
        content = join_inline_segments(buffer)
        if content:
            item.append(content)
            buffer.clear()
        item.append(segment["html"])
    content = join_inline_segments(buffer)
    if content:
        item.append(content)
    item.append("  </li>")
    return item


def render_paragraphs(doc, xml_paras, start, end, chapter, image_writer, rid_to_media, transformer):
    html_parts = []
    in_list = False

    def close_list():
        nonlocal in_list
        if in_list:
            html_parts.append("</ul>")
            in_list = False

    for idx in range(start, end):
        para = doc.paragraphs[idx]
        xml_para = xml_paras[idx]
        if is_toc(para):
            continue

        level = heading_level(para)
        text = clean_text(para.text)
        segments = paragraph_segments(xml_para, chapter, idx, para.text, image_writer, rid_to_media, transformer)

        if not segments:
            close_list()
            continue

        if level in (4, 5):
            close_list()
            tag = "h4" if level == 4 else "h5"
            html_parts.append(f"<{tag}>{inline_html(segments)}</{tag}>")
            html_parts.extend(segment["html"] for segment in segments if segment["kind"] == "block")
            continue

        if is_list_para(para):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            strip_list_marker_from_segments(segments)
            html_parts.extend(render_segments_as_list_item(segments))
            continue

        close_list()
        css = ' class="caption"' if para_style(para) == "Caption" else ""
        html_parts.extend(render_segments_as_blocks(segments, css))

    close_list()
    return html_parts


def collect_structure(doc):
    lnd = None
    refs = None
    chapters = []
    current_chapter = None
    section_count = 0
    subsection_count = 0

    for idx, para in enumerate(doc.paragraphs):
        if is_toc(para):
            continue
        text = clean_text(para.text)
        if not text:
            continue
        upper = text.upper()

        if "LỜI NÓI ĐẦU" in upper:
            lnd = {"start": idx, "end": None}
            continue
        if "TÀI LIỆU THAM KHẢO" in upper:
            refs = {"start": idx, "end": len(doc.paragraphs)}
            if current_chapter:
                current_chapter["end"] = idx
            continue

        level = heading_level(para)
        if level == 1 and upper in CHAPTER_HEADINGS:
            chapter_num = CHAPTER_HEADINGS[upper]
            if current_chapter:
                current_chapter["end"] = idx
            current_chapter = {
                "chapter": chapter_num,
                "title": text,
                "start": idx,
                "end": None,
                "sections": [],
            }
            chapters.append(current_chapter)
            section_count = 0
            subsection_count = 0
            if lnd and lnd["end"] is None:
                lnd["end"] = idx
            continue

        if not current_chapter:
            continue

        if level == 2:
            section_count += 1
            subsection_count = 0
            current_chapter["sections"].append({
                "level": 2,
                "section": section_count,
                "subsection": 0,
                "title": text,
                "start": idx,
                "end": None,
                "subsections": [],
            })
        elif level == 3 and current_chapter["sections"]:
            subsection_count += 1
            parent = current_chapter["sections"][-1]
            sub = {
                "level": 3,
                "section": parent["section"],
                "subsection": subsection_count,
                "title": text,
                "start": idx,
                "end": None,
            }
            parent["subsections"].append(sub)

    for chapter in chapters:
        if chapter["end"] is None:
            chapter["end"] = refs["start"] if refs else len(doc.paragraphs)
        entries = []
        for section in chapter["sections"]:
            entries.append(section)
            entries.extend(section["subsections"])
        for i, entry in enumerate(entries):
            entry["end"] = entries[i + 1]["start"] if i + 1 < len(entries) else chapter["end"]

    return {"lnd": lnd, "refs": refs, "chapters": sorted(chapters, key=lambda c: c["chapter"])}


def section_path(root, chapter, section, subsection=0):
    roman = ROMAN[section] if section < len(ROMAN) else str(section)
    filename = f"muc-{roman}.html" if subsection == 0 else f"muc-{roman}-{subsection}.html"
    return os.path.join(root, "chapters", f"ch{chapter}", filename)


def page_header(chapter, title, section_title=None):
    name = CHAPTER_NAMES[chapter]
    badge = f"ch{chapter}-hero"
    h2 = section_title or title
    parts = [
        f'<div class="sh2"><span class="badge {badge}">Chương {chapter} – {name}</span>',
        f"  <h2>{html.escape(h2)}</h2>",
        "</div>",
        '<div class="l3-content">',
    ]
    if section_title:
        parts.append(f'  <div class="l3-title">{html.escape(title)}</div>')
    return parts


def write_file(path, content, write):
    if not write:
        return
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)


def cleanup_generated(root, write):
    if not write:
        return
    for chapter in (1, 2, 3):
        chapter_dir = os.path.join(root, "chapters", f"ch{chapter}")
        os.makedirs(chapter_dir, exist_ok=True)
        for name in os.listdir(chapter_dir):
            if name == "index.html" or re.match(r"muc-[IVXLC]+(?:-\d+)?\.html$", name):
                os.remove(os.path.join(chapter_dir, name))


def render_chapter_index(doc, xml_paras, chapter, image_writer, rid_to_media, transformer):
    chapter_num = chapter["chapter"]
    intro_start = chapter["start"] + 1
    first_section = chapter["sections"][0]["start"] if chapter["sections"] else chapter["end"]
    content = [
        f'<div class="chapter-hero ch{chapter_num}-hero">',
        f'  <div class="cn">Chương {chapter_num}</div>',
        f"  <h2>{CHAPTER_NAMES[chapter_num].upper()}</h2>",
        "</div>",
    ]
    intro = render_paragraphs(doc, xml_paras, intro_start, first_section, chapter_num, image_writer, rid_to_media, transformer)
    if intro:
        content.append('<div class="l3-content">')
        content.extend(intro)
        content.append("</div>")
    content.extend([
        '<div class="ov-sec">',
        "  <h3>Nội dung</h3>",
        '  <ul class="cl">',
    ])
    for section in chapter["sections"]:
        roman = ROMAN[section["section"]]
        pid = f"ch{chapter_num}-{section['section']}"
        content.append(
            f'    <li><a href="#" onclick="loadPage(\'{pid}\');return false">{roman}. {html.escape(section["title"])}</a></li>'
        )
    content.extend(["  </ul>", "</div>"])
    return "\n".join(content)


def render_section_page(doc, xml_paras, chapter_num, section, image_writer, rid_to_media, transformer):
    parts = page_header(chapter_num, section["title"])
    body_end = section["subsections"][0]["start"] if section["subsections"] else section["end"]
    body = render_paragraphs(doc, xml_paras, section["start"] + 1, body_end, chapter_num, image_writer, rid_to_media, transformer)
    parts.extend(body)
    if section["subsections"]:
        parts.append('<div class="section-toc">')
        parts.append("  <h4>Nội dung</h4>")
        parts.append("  <ol>")
        for sub in section["subsections"]:
            pid = f"ch{chapter_num}-{section['section']}-{sub['subsection']}"
            parts.append(
                f'    <li><a href="#" onclick="loadPage(\'{pid}\');return false">{html.escape(sub["title"])}</a></li>'
            )
        parts.append("  </ol>")
        parts.append("</div>")
    parts.append("</div>")
    return "\n".join(parts)


def render_subsection_page(doc, xml_paras, chapter_num, section, sub, image_writer, rid_to_media, transformer):
    parts = page_header(chapter_num, sub["title"], section["title"])
    parts.extend(render_paragraphs(doc, xml_paras, sub["start"] + 1, sub["end"], chapter_num, image_writer, rid_to_media, transformer))
    parts.append("</div>")
    return "\n".join(parts)


def render_special_page(title, body):
    return "\n".join([f'<div class="sh2"><h2>{html.escape(title)}</h2></div>', '<div class="l3-content">', *body, "</div>"])


def front_matter_end(doc, lnd_start):
    end = lnd_start if lnd_start is not None else 0
    for idx, para in enumerate(doc.paragraphs[:end]):
        if clean_text(para.text).upper() == "MỤC LỤC":
            return idx
    return end


def extract(args):
    root = resolve_path(args.output)
    docx_path = resolve_path(args.input, root)
    if not os.path.exists(docx_path):
        raise SystemExit(f"Input DOCX not found: {docx_path}")

    doc = Document(docx_path)
    xml_paras, rid_to_media, media_blobs = package_data(docx_path)
    if len(xml_paras) != len(doc.paragraphs):
        raise SystemExit(f"Paragraph count mismatch: docx={len(doc.paragraphs)}, xml={len(xml_paras)}")

    structure = collect_structure(doc)
    if len(structure["chapters"]) != 3:
        raise SystemExit(f"Expected 3 chapters, found {len(structure['chapters'])}")

    transformer = load_omml_transformer()
    omml_count = count_omml(xml_paras)
    preflight_dependencies(media_blobs, omml_count, transformer, args.write)
    equation_mapping = load_equation_mapping(root)
    image_writer = ImageWriter(root, media_blobs, args.write, equation_mapping)
    cleanup_generated(root, args.write)
    image_writer.prepare()

    print("=== DOCX EXTRACT ===")
    print(f"Input: {docx_path}")
    print(f"Mode: {'WRITE' if args.write else 'DRY-RUN'}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Media blobs: {len(media_blobs)}")
    print(f"OMML math objects: {omml_count}")
    print(f"OMML transformer: {'yes' if transformer is not None else 'no'}")
    print(f"Reviewed equation mappings: {len(equation_mapping)}")

    manifest = {"input": docx_path, "chapters": []}

    for chapter in structure["chapters"]:
        chapter_num = chapter["chapter"]
        chapter_dir = os.path.join(root, "chapters", f"ch{chapter_num}")
        os.makedirs(chapter_dir, exist_ok=True)
        chapter_meta = {"chapter": chapter_num, "title": CHAPTER_NAMES[chapter_num], "sections": []}

        index_html = render_chapter_index(doc, xml_paras, chapter, image_writer, rid_to_media, transformer)
        write_file(os.path.join(chapter_dir, "index.html"), index_html, args.write)
        print(f"  ch{chapter_num}: {len(chapter['sections'])} sections")

        for section in chapter["sections"]:
            section_html = render_section_page(doc, xml_paras, chapter_num, section, image_writer, rid_to_media, transformer)
            write_file(section_path(root, chapter_num, section["section"]), section_html, args.write)
            section_meta = {
                "section": section["section"],
                "roman": ROMAN[section["section"]],
                "title": section["title"],
                "subsections": [],
            }
            for sub in section["subsections"]:
                sub_html = render_subsection_page(doc, xml_paras, chapter_num, section, sub, image_writer, rid_to_media, transformer)
                write_file(section_path(root, chapter_num, section["section"], sub["subsection"]), sub_html, args.write)
                section_meta["subsections"].append({"subsection": sub["subsection"], "title": sub["title"]})
                if "Câu hỏi ôn tập" in sub["title"]:
                    write_file(os.path.join(chapter_dir, "on-tap.html"), sub_html, args.write)
            chapter_meta["sections"].append(section_meta)
        manifest["chapters"].append(chapter_meta)

    if structure["lnd"]:
        front_end = front_matter_end(doc, structure["lnd"]["start"])
        body = render_paragraphs(doc, xml_paras, 0, front_end, 0, image_writer, rid_to_media, transformer)
        if body:
            write_file(
                os.path.join(root, "chapters", "tac-gia.html"),
                render_special_page("Tác giả và thông tin giáo trình", body),
                args.write,
            )

    if structure["lnd"]:
        body = render_paragraphs(
            doc, xml_paras, structure["lnd"]["start"] + 1, structure["lnd"]["end"], 0, image_writer, rid_to_media, transformer
        )
        write_file(os.path.join(root, "chapters", "loi-noi-dau.html"), render_special_page("Lời nói đầu", body), args.write)

    if structure["refs"]:
        body = render_paragraphs(
            doc, xml_paras, structure["refs"]["start"] + 1, structure["refs"]["end"], 0, image_writer, rid_to_media, transformer
        )
        write_file(os.path.join(root, "chapters", "tai-lieu-tham-khao.html"), render_special_page("Tài liệu tham khảo", body), args.write)

    image_writer.write_report()
    image_writer.write_equation_report(args.write or args.equation_report)
    if args.write:
        manifest_path = os.path.join(root, "tools", "docx_site_manifest.json")
        with open(manifest_path, "w", encoding="utf-8") as fh:
            json.dump(manifest, fh, ensure_ascii=False, indent=2)

    print(f"Images saved: {sum(len(v) for v in image_writer.saved.values())}")
    print(f"Image classes: {image_writer.class_counts}")
    print(f"Image conversion failures: {len(image_writer.failures)}")
    if image_writer.failures:
        print("Failed media:")
        for item in image_writer.failures[:20]:
            print(f"  ch{item['chapter']} {item['media']}: {item['reason']}")
        if len(image_writer.failures) > 20:
            print(f"  ... {len(image_writer.failures) - 20} more")
        if args.write:
            raise SystemExit("Image conversion failures detected; generated output is not complete.")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="CoHocLyThuyet_Full_New.docx")
    parser.add_argument("--output", default=os.getcwd())
    parser.add_argument("--write", action="store_true", help="Write generated fragments/images")
    parser.add_argument("--equation-report", action="store_true", help="Write tools/equation_report.json during dry-run")
    args = parser.parse_args()
    extract(args)


if __name__ == "__main__":
    main()
