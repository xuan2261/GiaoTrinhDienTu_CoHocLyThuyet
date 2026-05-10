"""
Script tạo file DeCuong_GTDT_CoHocLyThuyet.docx
Dựa theo format CHÍNH XÁC của DeCuong_GTDT_DoLuongVoTuyenDien.docx

Format rules (từ phân tích file gốc):
- KHÔNG dùng Title style → tất cả dùng Normal
- TẤT CẢ chữ MÀU ĐEN (không xanh, không color khác)
- Heading 1/2/3 cho navigation nhưng override font = Times New Roman 13pt ĐEN
- Bìa: Normal + CENTER, line spacing 1.2
- Nội dung: spacing before/after = 6pt, line spacing 1.3
- Lời nói đầu: first indent = 1.27cm
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

BLACK = RGBColor(0, 0, 0)

def force_black(element):
    """Xóa triệt để themeColor/themeTint khỏi XML, chỉ giữ val=000000.
    Áp dụng cho cả style element và run element."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for color_el in element.findall('.//w:color', nsmap):
        color_el.set(qn('w:val'), '000000')
        # Xóa theme attributes
        for attr in ['w:themeColor', 'w:themeTint', 'w:themeShade']:
            qattr = qn(attr)
            if qattr in color_el.attrib:
                del color_el.attrib[qattr]

doc = Document()

# ============ PAGE SETUP (khớp file gốc) ============
for section in doc.sections:
    section.page_width = Emu(7560945)
    section.page_height = Emu(10693400)
    section.left_margin = Emu(1260475)
    section.right_margin = Emu(540385)
    section.top_margin = Emu(900430)
    section.bottom_margin = Emu(720090)

# ============ STYLE OVERRIDES → đen, Times New Roman ============
# Override Title → force về Normal format, KHÔNG dùng style Title
title_s = doc.styles['Title']
title_s.font.name = 'Times New Roman'
title_s.font.size = Pt(13)
title_s.font.bold = False
title_s.font.color.rgb = BLACK
force_black(title_s.element)

# Override Heading 1/2/3 để text vẫn đen, Times New Roman
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[style_name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(13)
    s.font.bold = True
    s.font.color.rgb = BLACK
    s.paragraph_format.space_before = Emu(76200)
    s.paragraph_format.space_after = Emu(76200)
    s.paragraph_format.keep_with_next = None
    # XÓA themeColor khỏi XML style element
    force_black(s.element)

# Override Normal
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(13)
ns.font.color.rgb = BLACK
force_black(ns.element)

# ============ HELPER FUNCTIONS ============
def sf(run, bold=None, italic=None, underline=None, size=Pt(13)):
    """Set font on run: Times New Roman, ĐEN, size. Xóa triệt để themeColor."""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = size
    run.font.color.rgb = BLACK
    # Xóa themeColor khỏi run XML
    force_black(run._element)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if underline is not None: run.underline = underline

def cover_line(doc, text='', bold=False, italic=False, size=Pt(13)):
    """Bìa: Normal + CENTER, line spacing 1.2, không spacing"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic, size=size)
    return p

def h1(doc, text, center=True):
    """Heading 1 → cho navigation, font đen"""
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    sf(r, bold=True)
    return p

def h2(doc, text):
    """Heading 2 → mục I, II..., font đen"""
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    sf(r, bold=True)
    return p

def h3(doc, text):
    """Heading 3 → mục 1, 2..., font đen"""
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    sf(r, bold=True)
    return p

def sub_item(doc, text):
    """a), b)... → Normal, Bold+Italic, spacing 6pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Emu(76200)
    pf.space_after = Emu(76200)
    r = p.add_run(text)
    sf(r, bold=True, italic=True)

def content(doc, text, indent=True):
    """Paragraph nội dung: JUSTIFY, spacing 6pt, indent 1.27cm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Emu(76200)
    pf.space_after = Emu(76200)
    pf.line_spacing = 1.3
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if indent:
        pf.first_line_indent = Emu(457200)
    r = p.add_run(text)
    sf(r)

def review(doc, ch):
    """Ôn tập cuối chương"""
    h3(doc, f'Nội dung ôn tập chương {ch}')
    h3(doc, f'Câu hỏi trắc nghiệm chương {ch}')

# ================================================================
# TRANG BÌA 1 (Normal + CENTER, spacing 1.2)
# ================================================================
cover_line(doc, 'BỘ QUỐC PHÒNG', bold=True)
cover_line(doc, 'HỌC VIỆN HẢI QUÂN', bold=True)

for _ in range(11):
    cover_line(doc)

cover_line(doc, 'ĐỀ CƯƠNG CHI TIẾT', bold=True, size=Pt(14))
cover_line(doc)
cover_line(doc, 'GIÁO TRÌNH ĐIỆN TỬ', bold=True)
cover_line(doc, 'CƠ HỌC LÝ THUYẾT', bold=True, size=Pt(14))
cover_line(doc, '(Dùng cho đào tạo học viên sĩ quan cấp phân đội trình độ đại học)', italic=True)

for _ in range(7):
    cover_line(doc)

# TÁC GIẢ (JUSTIFY, như file gốc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('TÁC GIẢ:'); sf(r, bold=True)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Chủ biên: Đại tá, TS Nguyễn Lê Văn, CNK Khoa KTCS'); sf(r)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Tham gia biên soạn: 1. Thiếu tá, ThS Đinh Văn Tự, GV Khoa KTCS'); sf(r)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('                                 2. Đại úy, ThS Bùi Thanh Xuân, GV Khoa KTCS'); sf(r)

# ================================================================
# TRANG BÌA 2 (lặp lại như file gốc)
# ================================================================
doc.add_page_break()
cover_line(doc, 'ĐỀ CƯƠNG CHI TIẾT', bold=True, size=Pt(14))
cover_line(doc)
cover_line(doc, 'GIÁO TRÌNH ĐIỆN TỬ', bold=True)
cover_line(doc, 'CƠ HỌC LÝ THUYẾT', bold=True, size=Pt(14))
cover_line(doc, '(Dùng cho đào tạo học viên sĩ quan cấp phân đội trình độ đại học)', italic=True)

# ================================================================
# LỜI NÓI ĐẦU
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.2
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
r = p.add_run('LỜI NÓI ĐẦU'); sf(r, bold=True)

content(doc, 'Cơ học lý thuyết là một trong những môn khoa học cơ sở quan trọng nhất trong chương trình đào tạo kỹ sư, sĩ quan kỹ thuật tại các trường đại học nói chung và các học viện, nhà trường quân đội nói riêng.')
content(doc, 'Giáo trình được biên soạn nhằm đáp ứng yêu cầu giảng dạy và học tập cho học viên sĩ quan cấp phân đội trình độ đại học tại Học viện Hải quân, bám sát chương trình đào tạo đã được phê duyệt.')
content(doc, 'Giáo trình được cấu trúc thành 03 chương:')
content(doc, '- Chương 1: Tĩnh học')
content(doc, '- Chương 2: Động học')
content(doc, '- Chương 3: Động lực học', indent=False)
content(doc, 'Mỗi chương đều có tóm tắt lý thuyết, bài tập có hướng dẫn, bài tập tự làm, câu hỏi ôn tập và ôn tập trắc nghiệm.')
content(doc, 'Trong quá trình biên soạn, tập thể tác giả đã tham khảo nhiều tài liệu, giáo trình của các trường đại học trong và ngoài quân đội. Tuy đã cố gắng nhưng không tránh khỏi thiếu sót, rất mong nhận được sự góp ý của bạn đọc để giáo trình ngày càng hoàn thiện hơn.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format; pf.space_before = Emu(76200); pf.space_after = Emu(76200)
r = p.add_run('NHÓM TÁC GIẢ'); sf(r, bold=True)

for _ in range(10):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Emu(76200); pf.space_after = Emu(76200)

# ================================================================
# CHƯƠNG 1: TĨNH HỌC
# ================================================================
h1(doc, 'Chương 1')
h1(doc, 'TĨNH HỌC')

h2(doc, 'I. CÁC KHÁI NIỆM CƠ BẢN')
h3(doc, '1. Vật rắn tuyệt đối')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Ý nghĩa trong cơ học lý thuyết')
h3(doc, '2. Trạng thái cân bằng')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Phân loại cân bằng')
h3(doc, '3. Lực')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Các yếu tố của lực')
sub_item(doc, 'c) Đơn vị và biểu diễn lực')
h3(doc, '4. Mô men')
sub_item(doc, 'a) Mô men của lực đối với một điểm')
sub_item(doc, 'b) Mô men của lực đối với một trục')
sub_item(doc, 'c) Tính chất của mô men')
h3(doc, '5. Hệ lực')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Các loại hệ lực')
sub_item(doc, 'c) Hệ lực tương đương và hệ lực cân bằng')
h3(doc, '6. Ngẫu lực')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Mô men của ngẫu lực')
sub_item(doc, 'c) Tính chất của ngẫu lực')
h3(doc, '7. Vật tự do và vật không tự do')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Các loại liên kết')
h3(doc, '8. Lực liên kết, lực hoạt động, phản lực liên kết')
sub_item(doc, 'a) Lực liên kết')
sub_item(doc, 'b) Lực hoạt động')
sub_item(doc, 'c) Phản lực liên kết')

h2(doc, 'II. HỆ TIÊN ĐỀ TĨNH HỌC')
h3(doc, '1. Tiên đề 1: Cặp lực cân bằng')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ý nghĩa')
h3(doc, '2. Tiên đề 2: Thêm hay bớt cặp lực cân bằng')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Hệ quả: Trượt lực trên đường tác dụng')
h3(doc, '3. Tiên đề 3: Hình bình hành lực')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Quy tắc hình bình hành')
h3(doc, '4. Tiên đề 4: Lực tương tác')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ứng dụng')
h3(doc, '5. Tiên đề 5: Tiên đề hóa rắn')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ý nghĩa')
h3(doc, '6. Tiên đề 6: Giải phóng liên kết')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Phương pháp giải phóng liên kết')

h2(doc, 'III. MỘT SỐ PHẢN LỰC LIÊN KẾT THƯỜNG GẶP')
h3(doc, '1. Liên kết tựa')
sub_item(doc, 'a) Đặc điểm')
sub_item(doc, 'b) Phản lực liên kết')
h3(doc, '2. Liên kết dây mềm, thẳng')
sub_item(doc, 'a) Đặc điểm')
sub_item(doc, 'b) Phản lực liên kết')
h3(doc, '3. Liên kết bản lề')
sub_item(doc, 'a) Bản lề trụ')
sub_item(doc, 'b) Phản lực liên kết')
h3(doc, '4. Liên kết gối')
sub_item(doc, 'a) Gối cố định')
sub_item(doc, 'b) Gối di động')
sub_item(doc, 'c) Phản lực liên kết')
h3(doc, '5. Liên kết gối cầu')
sub_item(doc, 'a) Đặc điểm')
sub_item(doc, 'b) Phản lực liên kết')
h3(doc, '6. Liên kết ngàm')
sub_item(doc, 'a) Đặc điểm')
sub_item(doc, 'b) Phản lực liên kết')
h3(doc, '7. Liên kết thanh')
sub_item(doc, 'a) Đặc điểm')
sub_item(doc, 'b) Phản lực liên kết')

h2(doc, 'IV. ĐẶC TRƯNG CƠ BẢN CỦA HỆ LỰC')
h3(doc, '1. Véc tơ chính của hệ lực không gian')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Cách xác định')
h3(doc, '2. Mô men chính của hệ lực không gian với một tâm')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Công thức tính')
h3(doc, '3. Các dạng cơ bản của hệ lực không gian')
sub_item(doc, 'a) Hệ lực đồng quy')
sub_item(doc, 'b) Hệ lực song song')
sub_item(doc, 'c) Hệ ngẫu lực')
sub_item(doc, 'd) Hệ lực bất kỳ')

h2(doc, 'V. ĐIỀU KIỆN CÂN BẰNG')
h3(doc, '1. Điều kiện và phương trình cân bằng của hệ lực bất kỳ trong không gian')
sub_item(doc, 'a) Điều kiện cần và đủ')
sub_item(doc, 'b) Hệ phương trình cân bằng')
h3(doc, '2. Điều kiện và phương trình cân bằng của hệ lực đặc biệt')
sub_item(doc, 'a) Hệ lực đồng quy')
sub_item(doc, 'b) Hệ lực song song')
sub_item(doc, 'c) Hệ lực phẳng')
sub_item(doc, 'd) Hệ lực phẳng đồng quy')
sub_item(doc, 'e) Hệ lực phẳng song song')

h2(doc, 'VI. BÀI TẬP')
h3(doc, '1. Phương pháp giải bài toán cân bằng vật rắn')
sub_item(doc, 'a) Các bước giải')
sub_item(doc, 'b) Ví dụ minh họa')
h3(doc, '2. Bài tập có hướng dẫn')
h3(doc, '3. Bài tập tự làm')

review(doc, 1)

# ================================================================
# CHƯƠNG 2: ĐỘNG HỌC
# ================================================================
h1(doc, 'Chương 2')
h1(doc, 'ĐỘNG HỌC')

h2(doc, 'I. CHUYỂN ĐỘNG CỦA CHẤT ĐIỂM')
h3(doc, '1. Khảo sát chuyển động bằng véc tơ')
sub_item(doc, 'a) Phương trình chuyển động')
sub_item(doc, 'b) Vận tốc')
sub_item(doc, 'c) Gia tốc')
h3(doc, '2. Khảo sát chuyển động bằng tọa độ Đề các')
sub_item(doc, 'a) Phương trình chuyển động')
sub_item(doc, 'b) Vận tốc')
sub_item(doc, 'c) Gia tốc')
h3(doc, '3. Khảo sát chuyển động bằng tọa độ tự nhiên')
sub_item(doc, 'a) Hệ trục tọa độ tự nhiên')
sub_item(doc, 'b) Vận tốc')
sub_item(doc, 'c) Gia tốc tiếp tuyến và gia tốc pháp tuyến')

h2(doc, 'II. CHUYỂN ĐỘNG CỦA VẬT RẮN')
h3(doc, '1. Chuyển động tịnh tiến của vật rắn')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Tính chất')
sub_item(doc, 'c) Phương trình chuyển động')
h3(doc, '2. Chuyển động quay quanh một trục cố định')
sub_item(doc, 'a) Định nghĩa và phương trình chuyển động')
sub_item(doc, 'b) Vận tốc góc và gia tốc góc')
sub_item(doc, 'c) Vận tốc và gia tốc của điểm thuộc vật')

h2(doc, 'III. HỢP CHUYỂN ĐỘNG ĐIỂM')
h3(doc, '1. Định nghĩa và các loại chuyển động')
sub_item(doc, 'a) Chuyển động tuyệt đối')
sub_item(doc, 'b) Chuyển động tương đối')
sub_item(doc, 'c) Chuyển động kéo theo')
h3(doc, '2. Các định lý về hợp vận tốc và hợp gia tốc')
sub_item(doc, 'a) Định lý hợp vận tốc')
sub_item(doc, 'b) Định lý hợp gia tốc (Coriolis)')
sub_item(doc, 'c) Gia tốc Coriolis')

h2(doc, 'IV. CHUYỂN ĐỘNG SONG PHẲNG')
h3(doc, '1. Khái niệm và mô hình chuyển động song phẳng')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Phân tích chuyển động')
h3(doc, '2. Khảo sát chuyển động của vật rắn')
sub_item(doc, 'a) Phương trình chuyển động')
sub_item(doc, 'b) Vận tốc góc và gia tốc góc')
h3(doc, '3. Khảo sát chuyển động của các điểm thuộc vật')
sub_item(doc, 'a) Phân bố vận tốc')
sub_item(doc, 'b) Tâm vận tốc tức thời')
sub_item(doc, 'c) Phân bố gia tốc')

h2(doc, 'V. BÀI TẬP ĐỘNG HỌC')
h3(doc, '1. Tóm tắt lý thuyết')
h3(doc, '2. Bài tập có hướng dẫn')
h3(doc, '3. Bài tập tự làm')

review(doc, 2)

# ================================================================
# CHƯƠNG 3: ĐỘNG LỰC HỌC
# ================================================================
h1(doc, 'Chương 3')
h1(doc, 'ĐỘNG LỰC HỌC')

h2(doc, 'I. CÁC KHÁI NIỆM CƠ BẢN')
h3(doc, '1. Vật thể')
sub_item(doc, 'a) Chất điểm')
sub_item(doc, 'b) Cơ hệ')
h3(doc, '2. Lực')
sub_item(doc, 'a) Lực nội và lực ngoại')
sub_item(doc, 'b) Tính chất của nội lực')
h3(doc, '3. Hệ quy chiếu quán tính')
sub_item(doc, 'a) Định nghĩa')
sub_item(doc, 'b) Ý nghĩa trong động lực học')

h2(doc, 'II. CÁC ĐỊNH LUẬT CƠ BẢN')
h3(doc, '1. Tiên đề 1 – Định luật quán tính')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ý nghĩa')
h3(doc, '2. Tiên đề 2 – Định luật cơ bản của động lực học')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Các hệ quả')
h3(doc, '3. Tiên đề 3 – Định luật tác dụng và phản tác dụng')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ý nghĩa')
h3(doc, '4. Tiên đề 4 – Định luật tính độc lập giữa tác dụng của lực')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ứng dụng')
h3(doc, '5. Tiên đề 5 – Tiên đề giải phóng liên kết')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Ứng dụng trong động lực học')

h2(doc, 'III. PHƯƠNG TRÌNH VI PHÂN CHUYỂN ĐỘNG CỦA CHẤT ĐIỂM VÀ CƠ HỆ')
h3(doc, '1. Dạng véc tơ')
sub_item(doc, 'a) Phương trình vi phân chuyển động của chất điểm')
sub_item(doc, 'b) Phương trình vi phân chuyển động của cơ hệ')
h3(doc, '2. Dạng tọa độ Đề các')
sub_item(doc, 'a) Hệ phương trình cho chất điểm')
sub_item(doc, 'b) Hệ phương trình cho cơ hệ')

h2(doc, 'IV. HAI BÀI TOÁN CƠ BẢN CỦA ĐỘNG LỰC HỌC')
h3(doc, '1. Bài toán thuận')
sub_item(doc, 'a) Phát biểu bài toán')
sub_item(doc, 'b) Phương pháp giải')
sub_item(doc, 'c) Ví dụ minh họa')
h3(doc, '2. Bài toán ngược')
sub_item(doc, 'a) Phát biểu bài toán')
sub_item(doc, 'b) Phương pháp giải')
sub_item(doc, 'c) Ví dụ minh họa')

h2(doc, 'V. CÁC ĐỊNH LÝ TỔNG QUÁT')
h3(doc, '1. Định lý chuyển động khối tâm')
sub_item(doc, 'a) Phát biểu')
sub_item(doc, 'b) Hệ quả: Bảo toàn chuyển động khối tâm')
h3(doc, '2. Định lý động lượng')
sub_item(doc, 'a) Động lượng của chất điểm và cơ hệ')
sub_item(doc, 'b) Phát biểu định lý')
sub_item(doc, 'c) Hệ quả: Bảo toàn động lượng')
h3(doc, '3. Định lý mô men động lượng')
sub_item(doc, 'a) Mô men động lượng')
sub_item(doc, 'b) Phát biểu định lý')
sub_item(doc, 'c) Hệ quả: Bảo toàn mô men động lượng')
h3(doc, '4. Định lý động năng')
sub_item(doc, 'a) Công của lực')
sub_item(doc, 'b) Động năng')
sub_item(doc, 'c) Phát biểu định lý')
sub_item(doc, 'd) Hệ quả: Bảo toàn cơ năng')

h2(doc, 'VI. LÝ THUYẾT VA CHẠM')
h3(doc, '1. Các đặc điểm và giả thiết về va chạm')
sub_item(doc, 'a) Định nghĩa va chạm')
sub_item(doc, 'b) Các giả thiết cơ bản')
sub_item(doc, 'c) Hệ số phục hồi')
h3(doc, '2. Các định lý tổng quát áp dụng vào va chạm')
sub_item(doc, 'a) Định lý động lượng trong va chạm')
sub_item(doc, 'b) Định lý mô men động lượng trong va chạm')
h3(doc, '3. Hai bài toán cơ bản về va chạm')
sub_item(doc, 'a) Va chạm thẳng xuyên tâm')
sub_item(doc, 'b) Va chạm xiên')

h2(doc, 'VII. BÀI TẬP')
h3(doc, '1. Tóm tắt lý thuyết')
h3(doc, '2. Bài tập có hướng dẫn')
h3(doc, '3. Bài tập tự làm')

review(doc, 3)

# ================================================================
# TÀI LIỆU THAM KHẢO
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format; pf.space_before = Emu(76200); pf.space_after = Emu(76200)
r = p.add_run('TÀI LIỆU THAM KHẢO'); sf(r, bold=True)

for ref_text, ref_title, ref_pub in [
    ('1. Nguyễn Phong Điền, ', 'Bài tập Cơ học kỹ thuật', ', NXB Bách khoa Hà Nội, 2023.'),
    ('2. Nguyễn Văn Khang, ', 'Cơ học kỹ thuật', ', NXB Bách khoa Hà Nội, 2024.'),
    ('3. Đỗ Sanh, ', 'Cơ học kỹ thuật tập 1, 2', ', NXB Bách khoa Hà Nội, 2024.'),
    ('4. Đỗ Sanh, ', 'Bài tập Cơ học kỹ thuật tập 1, 2', ', NXB Bách khoa Hà Nội, 2025.'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(ref_text); sf(r)
    r = p.add_run(ref_title); sf(r, italic=True)
    r = p.add_run(ref_pub); sf(r)

# ================================================================
# PHỤ LỤC
# ================================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHỤ LỤC'); sf(r, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('QUY CÁCH TRÌNH BÀY GIÁO TRÌNH ĐIỆN TỬ'); sf(r, bold=True, underline=True)

content(doc, 'Giáo trình "Cơ học lý thuyết" được biên soạn theo quy cách của một giáo trình điện tử. Nội dung giáo trình được chia thành các chương rõ ràng, với mục lục bên trái giúp người học dễ dàng tra cứu và truy cập nhanh đến từng phần nội dung.')
content(doc, 'Phương pháp trình bày giáo trình điện tử theo phương pháp này không chỉ thuận tiện cho người học khi học online, mà còn thích hợp cho việc bổ sung, tích hợp với các phần mềm mô phỏng, bài tập trắc nghiệm tương tác và các tài nguyên đa phương tiện khác.')
content(doc, 'Sau quá trình nghiên cứu, nhóm tác giả đã biên soạn giáo trình trên cơ sở phần mềm soạn thảo Giáo trình điện tử eXeLearning – là phần mềm soạn giáo trình điện tử nổi tiếng, được sử dụng thông dụng ở nhiều trường đại học trên thế giới.')
content(doc, 'Phần mềm eXeLearning hỗ trợ soạn thảo nội dung đa phương tiện, tích hợp các đối tượng tương tác như câu hỏi trắc nghiệm, liên kết, hình ảnh, video. Đặc biệt, giáo trình có thể xuất ra định dạng HTML, SCORM, IMS, phù hợp với nhiều nền tảng học trực tuyến.')
content(doc, 'Giáo trình cũng được biên soạn thành 01 Giáo trình bản in, dựa theo Hướng dẫn của BTTM Cục Quân huấn – Nhà trường về Quy trình biên soạn, quy cách cho Giáo trình điện tử tại các nhà trường quân đội.')

# ============ SAVE ============
output = 'DeCuong_GTDT_CoHocLyThuyet.docx'
doc.save(output)
print(f'OK: {output}')
